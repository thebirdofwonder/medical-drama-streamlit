"""
医学論文PDFまたは台本 → 台本用意 → AIレビュー → VOICEVOX音声 → 静止画背景 → MP4 生成
Streamlit アプリ（macOS / Apple Silicon 向け）
"""

from __future__ import annotations

import io
import json
import os
import re
import shutil
import struct
import tempfile
import wave
from datetime import datetime
from pathlib import Path
from typing import Any, Callable

import requests
import streamlit as st
from docx import Document
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader

# ---------------------------------------------------------------------------
# 定数
# ---------------------------------------------------------------------------
VOICEVOX_URL = "http://127.0.0.1:50021"
# 初期値（UIで変更可）。声優「青山龍星」のノーマル
DEFAULT_SPEAKER_ID = 13
DEFAULT_SPEAKER_NAME = "青山龍星"
DEFAULT_STYLE_NAME = "ノーマル"
# 読み上げ速度（画面のスライダーで変更可）
VOICEVOX_SPEED_SCALE = 1.0  # 初期値
VOICEVOX_SPEED_MIN = 0.8
VOICEVOX_SPEED_MAX = 1.5
VOICEVOX_SPEED_STEP = 0.1
VIDEO_SIZE = (1920, 1080)
CREDIT_TEXT = f"音声\nVOICEVOX：{DEFAULT_SPEAKER_NAME}"
# 後方互換（古いコード参照用）
SPEAKER_ID = DEFAULT_SPEAKER_ID
DISCLAIMER_TEXT = (
    "本動画は、公開された症例報告を参考に制作した医学教育用フィクションです。"
    "人物、会話、状況設定は創作であり、ナレーションには合成音声を使用しています。"
)
# エンディング固定文（雛形どおり）
ENDING_FICTION_NOTICE = (
    "本動画は医学教育を目的としたフィクションです。"
    "登場人物、氏名、年齢、職業、会話、診療場面、時系列、"
    "および検査値・状況設定は実在のものではありません。"
)
ENDING_FOOTER = "詳細な出典・ライセンスは動画説明欄に記載しています。"
DEFAULT_REFERENCE_EXAMPLE = (
    "Lim J, Wenham T. An Atypical Presentation of Mycoplasma pneumoniae "
    "Infection Mimicking Acute Surgical Abdomen in an Adult. "
    "Cureus. 2024;16(11):e73665. doi:10.7759/cureus.73665"
)
SCENE_INTERVAL_SEC = 60.0  # 1分ごとに背景切替
ENDING_DURATION_SEC = 10.0  # エンディング画面を出し続ける秒数（フェード完了後）
ENDING_FADE_SEC = 5.0  # 本編音声終了後、エンディングへ完全移行するフェード秒数
INTRO_SILENCE_SEC = 3.0  # 冒頭の無音（朗読開始までの余白）
BGM_FILENAME = "bgm.mp3"
# Pixabay のフリー音源（暗め・シリアス寄り）。取得失敗時は簡易BGMを自動生成します。
BGM_CANDIDATE_URLS = [
    "https://cdn.pixabay.com/download/audio/2022/01/18/audio_d0c6ff1bab.mp3?filename=dark-ambient-11495.mp3",
    "https://cdn.pixabay.com/download/audio/2021/08/09/audio_dc39bde808.mp3?filename=cinematic-documentary-11521.mp3",
]
WORK_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = WORK_DIR / "outputs"
# ルビ辞書ファイル（用語とよみの対照表）
RUBY_DICT_PATH = WORK_DIR / "data" / "medical_ruby_dict.tsv"
# 修正反映後のダウンロード用辞書
RUBY_DICT_EXPORT_NAME = "ルビ辞書.txt"
RUBY_DICT_EXPORT_PATH = OUTPUT_DIR / RUBY_DICT_EXPORT_NAME

# 医療関連の著作権フリー背景（Unsplash）。旧・風景キャッシュは使わない
MEDICAL_BG_DIR = OUTPUT_DIR / "medical_backgrounds"
LANDSCAPE_DIR = MEDICAL_BG_DIR  # 互換エイリアス
# 参考文献の前回入力（outputs/ は GitHub に上がらない）
REFERENCE_SAVE_PATH = OUTPUT_DIR / "last_reference.txt"


def get_desktop_dir() -> Path:
    """macOS のデスクトップフォルダ（英語名 / 日本語名の両方を探す）。"""
    for name in ("Desktop", "デスクトップ"):
        p = Path.home() / name
        if p.is_dir():
            return p
    p = Path.home() / "Desktop"
    p.mkdir(parents=True, exist_ok=True)
    return p


def make_desktop_mp4_filename(title: str = "") -> str:
    """デスクトップ保存用のファイル名（上書きしにくいよう日時つき）。"""
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{make_title_basename(title)}_{stamp}.mp4"


def make_title_basename(title: str = "", fallback: str = "medical_drama") -> str:
    """タイトルをファイル名に使える文字だけにする。"""
    raw = (title or "").strip() or fallback
    safe = re.sub(r'[\\/:*?"<>|\s]+', "_", raw)
    safe = re.sub(r"_+", "_", safe).strip("._")[:80] or fallback
    return safe


def make_script_docx_filename(title: str = "") -> str:
    """台本 Word 用ファイル名（動画タイトルと同一）。"""
    return f"{make_title_basename(title)}.docx"


def load_saved_reference_text() -> str:
    """前回保存した参考文献を読み込む。"""
    try:
        if REFERENCE_SAVE_PATH.is_file():
            return REFERENCE_SAVE_PATH.read_text(encoding="utf-8").strip()
    except OSError:
        pass
    return ""


def save_reference_text(text: str) -> None:
    """参考文献をローカルに保存（次回起動時も残す・上書き可）。"""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    body = (text or "").rstrip()
    REFERENCE_SAVE_PATH.write_text(
        (body + "\n") if body else "",
        encoding="utf-8",
    )


def persist_reference_from_widget() -> None:
    """参考文献テキスト欄の変更をファイルへ保存する（Streamlit on_change用）。"""
    save_reference_text(str(st.session_state.get("reference_text") or ""))
# Unsplash（著作権フリー）の医療関連写真。病院・手術室・病室・検査機器など。
# 出典は動画説明欄への記載を推奨（Unsplash License）。
MEDICAL_BACKGROUND_URLS = [
    # 病院廊下
    "https://images.unsplash.com/photo-1519494026892-80bbd2d6fd0d?auto=format&fit=crop&w=1920&h=1080&q=80",
    # 手術室・執刀
    "https://images.unsplash.com/photo-1551076805-e1869033e561?auto=format&fit=crop&w=1920&h=1080&q=80",
    # モニター・検査機器
    "https://images.unsplash.com/photo-1579684385127-1ef15d508118?auto=format&fit=crop&w=1920&h=1080&q=80",
    # 病室・ベッド
    "https://images.unsplash.com/photo-1516549655169-df83a0774514?auto=format&fit=crop&w=1920&h=1080&q=80",
    # 聴診器など診療用具
    "https://images.unsplash.com/photo-1584982751601-97dcc096659c?auto=format&fit=crop&w=1920&h=1080&q=80",
    # 手術灯・オペ室
    "https://images.unsplash.com/photo-1551601651-2a8555f1a136?auto=format&fit=crop&w=1920&h=1080&q=80",
    # 手術用具
    "https://images.unsplash.com/photo-1581594693702-fbdc51b2763b?auto=format&fit=crop&w=1920&h=1080&q=80",
    # レントゲン・画像検査
    "https://images.unsplash.com/photo-1530497610245-94d3c16cda28?auto=format&fit=crop&w=1920&h=1080&q=80",
    # 医療スタッフ・診療
    "https://images.unsplash.com/photo-1576091160399-112ba8d25d1d?auto=format&fit=crop&w=1920&h=1080&q=80",
    # 病院内の医療機器
    "https://images.unsplash.com/photo-1666214280557-f1b5022eb634?auto=format&fit=crop&w=1920&h=1080&q=80",
    # 救急・医療現場
    "https://images.unsplash.com/photo-1505751172876-fa1923c5c528?auto=format&fit=crop&w=1920&h=1080&q=80",
    # 検査・ラボ
    "https://images.unsplash.com/photo-1582719471384-894fbb16e074?auto=format&fit=crop&w=1920&h=1080&q=80",
]
# 台本の場面に近い写真を優先するための対応（URLリストの番号）
THEME_TO_BG_INDEX = {
    "er": 10,
    "icu": 2,
    "surgery": 1,
    "lab": 11,
    "ward": 3,
    "consult": 8,
    "pharma": 4,
    "ambulance": 10,
}
LANDSCAPE_IMAGE_URLS = MEDICAL_BACKGROUND_URLS  # 互換エイリアス
# 1区間＝1字幕にするため短め（句点で区切り、長文のみここで切る）
MAX_VOICEVOX_CHARS = 90
# 字幕切替の時間精度（低いと最大で約 1/fps 秒ずれる）
SUBTITLE_VIDEO_FPS = 8
DEFAULT_FOOTNOTE = ""
# Anthropic の現行モデル（旧 claude-sonnet-4-20250514 は引退済み）
CLAUDE_MODEL_CANDIDATES = [
    "claude-sonnet-4-6",
    "claude-sonnet-5",
    "claude-sonnet-4-5",
    "claude-haiku-4-5",
]
# レビュー用に送る台本の上限（長すぎると API が失敗しやすい）
REVIEW_SCRIPT_MAX_CHARS = 12000
# 論文PDF→台本化：論文本文の送付上限・VOICEVOX 1倍速 10〜12分目安
PAPER_TEXT_MAX_CHARS = 100000
# 日本語ナレーション目安 約300〜330字/分 × 10〜12分（VOICEVOX 1.0倍速）
DRAMA_SCRIPT_TARGET_CHARS_MIN = 3000
DRAMA_SCRIPT_TARGET_CHARS_MAX = 4000


# ---------------------------------------------------------------------------
# ユーティリティ: テキスト／PDF読み込み
# ---------------------------------------------------------------------------
def extract_text_from_bytes(name: str, raw: bytes) -> str:
    """ファイル名と中身（バイト列）から台本テキストを取り出す。"""
    fname = (name or "").lower().strip()
    if isinstance(raw, str):
        raw = raw.encode("utf-8")
    if not isinstance(raw, (bytes, bytearray)):
        raw = bytes(raw or b"")

    if fname.endswith(".txt") or fname.endswith(".text") or fname.endswith(".md"):
        for encoding in ("utf-8", "utf-8-sig", "cp932", "shift_jis"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace")

    if fname.endswith(".docx"):
        doc = Document(io.BytesIO(raw))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # 表の中の文章も拾う
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
        return "\n".join(parts)

    if fname.endswith(".pdf"):
        return extract_text_from_pdf_bytes(raw)

    raise ValueError(
        f"対応形式は .txt / .docx / .pdf です（受け取ったファイル: {name or '不明'}）。"
    )


def extract_text_from_pdf_bytes(raw: bytes) -> str:
    """PDF（医学論文など）から文字を取り出す。"""
    if not isinstance(raw, (bytes, bytearray)):
        raw = bytes(raw or b"")
    if not raw:
        raise ValueError("PDFが空です。")
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as e:  # noqa: BLE001
        raise ValueError(f"PDFを開けませんでした: {e}") from e
    pages: list[str] = []
    for page in reader.pages:
        try:
            t = (page.extract_text() or "").strip()
        except Exception:
            t = ""
        if t:
            pages.append(t)
    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError(
            "PDFから文字を取り出せませんでした。"
            "スキャン画像のみのPDFの可能性があります。"
        )
    return text


def commit_loaded_script(text: str, source_id: str) -> None:
    """読み込んだ台本をセッションに入れ、以降の工程を最初からにする（ルビなし）。"""
    clear_review_decision_widgets(st.session_state.get("review"))
    plain = strip_voicevox_ruby(text or "").strip()
    st.session_state.raw_script = plain
    st.session_state.final_script = plain
    st.session_state.final_script_editor = plain
    st.session_state.final_script_editor_widget = plain
    st.session_state.review = None
    st.session_state.review_done = False
    st.session_state.skip_review = False
    st.session_state.script_confirmed = False
    st.session_state.ruby_ready = False
    st.session_state.ruby_script = ""
    st.session_state.ruby_script_baseline = ""
    st.session_state.ruby_skipped = False
    st.session_state.pop("ruby_script_editor", None)
    st.session_state.mp4_bytes = None
    st.session_state.mp4_path = ""
    st.session_state.review_apply_log = []
    st.session_state.review_manual_log = []
    st.session_state._script_file_id = source_id


# ---------------------------------------------------------------------------
# AI レビュー（Anthropic Claude API / フォールバック簡易レビュー）
# ---------------------------------------------------------------------------
ENV_FILE = Path(__file__).resolve().parent / ".env"


def load_dotenv_file(path: Path | None = None) -> None:
    """
    .env から KEY=VALUE を読み、未設定の環境変数だけ入れる。
    （GitHub には上げないローカル専用ファイル）
    """
    env_path = path or ENV_FILE
    if not env_path.is_file():
        return
    try:
        raw = env_path.read_text(encoding="utf-8")
    except OSError:
        return
    for line in raw.splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, _, val = line.partition("=")
        key = key.strip()
        val = val.strip().strip('"').strip("'")
        if not key:
            continue
        # すでにターミナル等で設定済みなら上書きしない
        if os.environ.get(key, "").strip():
            continue
        os.environ[key] = val


def save_api_key_to_env_file(api_key: str, path: Path | None = None) -> Path:
    """
    APIキーを .env に保存（または更新）。他の行はそのまま残す。
    """
    env_path = path or ENV_FILE
    key_name = "ANTHROPIC_API_KEY"
    api_key = (api_key or "").strip()
    if not api_key:
        raise ValueError("保存する APIキーが空です。")

    lines: list[str] = []
    found = False
    if env_path.is_file():
        try:
            old = env_path.read_text(encoding="utf-8").splitlines()
        except OSError:
            old = []
        for line in old:
            stripped = line.strip()
            if stripped.startswith("#") or "=" not in stripped:
                lines.append(line)
                continue
            k, _, _ = stripped.partition("=")
            if k.strip() == key_name:
                lines.append(f"{key_name}={api_key}")
                found = True
            else:
                lines.append(line)
    if not found:
        if lines and lines[-1].strip():
            lines.append("")
        lines.append(f"# Anthropic Claude API（このファイルは GitHub に上げません）")
        lines.append(f"{key_name}={api_key}")

    env_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    # すぐに使えるように環境変数にも入れる
    os.environ[key_name] = api_key
    return env_path


def get_api_key() -> str:
    """
    APIキー取得の優先順位:
    1. 環境変数 ANTHROPIC_API_KEY（.env 読込後含む）
    2. Streamlit secrets（.streamlit/secrets.toml）
    コードには直書きしない。
    """
    load_dotenv_file()
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if key:
        return key
    try:
        return str(st.secrets.get("ANTHROPIC_API_KEY", "") or "").strip()
    except Exception:
        return ""


def build_review_prompt(script: str) -> str:
    return f"""あなたは現役の救急・集中治療に詳しい医療監修医です。
あわせて VOICEVOX（音声合成）向けの読み付けも行います。
以下の医学ドラマ台本を検証し、JSONオブジェクトだけを返してください。
前後に説明文・Markdown・コードフェンスは付けないでください。

厳守ルール:
- 有効なJSONのみ（末尾カンマ禁止）
- 文字列内に半角ダブルクォート " を書かない（必要なら『』や「」を使う）
- 各指摘配列は最大5件
- original は40文字以内、issue/suggestion は120文字以内
- suggestion は「台本にそのまま差し替える完成文」だけを書く
  （禁止例: 「編集メモを削除する」「確定文にして」「地の文として採用」などの作業手順・解説）

【最重要・禁止事項（レビュー指摘について）】
この台本では、音声合成の誤読を防ぐため、医学用語・専門用語を意図的にカタカナ表記しています。
したがって次は絶対に指摘しないでください（medical_contradictions / awkward_for_doctors / immersion_improvements のいずれにも含めない）:
- 医学用語・専門用語がカタカナであること
- カタカナを漢字に直す提案
- 「医師なら漢字で書く」「カタカナは不自然」といった表記スタイルの指摘
内容の医学的正しさ・現場表現・臨場感のみを見てください。

【VOICEVOXルビ付与（必須）】
読み間違えやすい漢字の医学用語・専門用語・難読語について、ruby_annotations に列挙してください。
- surface: 台本中の表記そのもの（漢字など。すでにカタカナだけの語は原則不要）
- reading: 正しい読み（ひらがな、またはカタカナ）
- すでに {{表記|よみ}} 形式のものがある場合は重複させない
- 人名の難読、薬品名、疾患名、手技名、略語の読みなどを優先（最大40件）
- 読みは実際の医療現場の読みに合わせる（例: 心筋梗塞→しんきんこうそく）

観点:
1. medical_contradictions … 医学的に矛盾している箇所
2. awkward_for_doctors … 現役医師が聞くと違和感がある表現（カタカナ表記そのものは対象外）
3. immersion_improvements … 臨場感が増す修正（カタカナを漢字にする案は出さない）
4. ruby_annotations … VOICEVOX用ルビ一覧

形式:
{{
  "medical_contradictions": [
    {{"original": "引用", "issue": "問題", "suggestion": "差し替え用の完成文のみ"}}
  ],
  "awkward_for_doctors": [
    {{"original": "引用", "issue": "問題", "suggestion": "差し替え用の完成文のみ"}}
  ],
  "immersion_improvements": [
    {{"original": "引用", "issue": "問題", "suggestion": "差し替え用の完成文のみ"}}
  ],
  "ruby_annotations": [
    {{"surface": "心筋梗塞", "reading": "しんきんこうそく"}}
  ]
}}

該当が無い観点は空配列 [] にしてください。

台本:
---
{script}
---
"""


def normalize_voicevox_reading(reading: str) -> str:
    """ルビの読みから余計な記号を除く。"""
    reading = (reading or "").strip()
    reading = reading.replace(" ", "").replace("　", "")
    reading = reading.replace("{", "").replace("}", "").replace("｛", "").replace("｝", "")
    if "|" in reading:
        reading = reading.split("|")[-1]
    if "｜" in reading:
        reading = reading.split("｜")[-1]
    return reading.strip()


# ルビ区切り: 半角 { } | と全角 ｛ ｝ ｜ は同じものとして扱う
_RUBY_OPEN = r"[{｛]"
_RUBY_PIPE = r"[|｜]"
_RUBY_CLOSE = r"[}｝]"
_RUBY_TAG_RE = re.compile(
    _RUBY_OPEN + r"([^|｜\n]+)" + _RUBY_PIPE + r"([^}｝\n]+)" + _RUBY_CLOSE
)


def canonicalize_voicevox_ruby_delimiters(text: str) -> str:
    """
    ルビの {｝| と ｛｝｜ を区別せず、すべて半角 {表記|よみ} にそろえる。
    """
    if not text:
        return ""

    def _repl(match: re.Match) -> str:
        surface = match.group(1)
        reading = normalize_voicevox_reading(match.group(2))
        return "{" + surface + "|" + reading + "}"

    return _RUBY_TAG_RE.sub(_repl, text)


def apply_voicevox_ruby(
    script: str,
    annotations: list[dict[str, str]],
    *,
    fullwidth: bool = True,
) -> str:
    """
    VOICEVOXルビを台本へ付与する。
    fullwidth=True のとき ｛表記｜よみ｝、False のとき {表記|よみ}。
    半角/全角の区切り記号は区別しない。既にあるルビは壊さない。

    同じ位置に複数の辞書語が当てはまるときは、文字数が多い用語を優先する。
    例: 「所」と「所見」なら「所見」を採用する。
    """
    text = canonicalize_voicevox_ruby_delimiters(script or "")
    if not annotations:
        return (
            to_fullwidth_ruby_delimiters(text)
            if fullwidth
            else text
        )

    protected: dict[str, str] = {}

    def _protect(match: re.Match) -> str:
        key = f"\x00RUBY{len(protected)}\x00"
        protected[key] = match.group(0)
        return key

    def protect_existing(src: str) -> str:
        return _RUBY_TAG_RE.sub(_protect, src)

    pairs: list[tuple[str, str]] = []
    seen: set[str] = set()
    for item in annotations:
        surface = str(item.get("surface") or "").strip()
        reading = normalize_voicevox_reading(str(item.get("reading") or ""))
        if not surface or not reading:
            continue
        if surface in seen:
            continue
        if surface == reading:
            continue
        if len(reading) < 1:
            continue
        seen.add(surface)
        pairs.append((surface, reading))

    # 長い用語を先に試し、短い用語（例:「所」）より「所見」を優先
    pairs.sort(key=lambda x: (len(x[0]), x[0]), reverse=True)
    text = protect_existing(text)

    # プレースホルダ以外の平文だけを、左から最長一致で置換する
    placeholder_re = re.compile(r"\x00RUBY\d+\x00")
    pieces: list[str] = []
    last = 0
    for m in placeholder_re.finditer(text):
        if m.start() > last:
            pieces.append(
                _apply_ruby_longest_match(text[last:m.start()], pairs, fullwidth)
            )
        pieces.append(m.group(0))
        last = m.end()
    if last < len(text):
        pieces.append(_apply_ruby_longest_match(text[last:], pairs, fullwidth))
    text = "".join(pieces)

    for key, val in protected.items():
        text = text.replace(key, val)
    text = canonicalize_voicevox_ruby_delimiters(text)
    return to_fullwidth_ruby_delimiters(text) if fullwidth else text


def _apply_ruby_longest_match(
    segment: str,
    pairs: list[tuple[str, str]],
    fullwidth: bool,
) -> str:
    """
    平文の先頭から順に見て、その位置で一致する用語のうち
    いちばん文字数が多いものをルビにする。
    pairs は文字数の多い順に並べておくこと。
    """
    if not segment or not pairs:
        return segment
    out: list[str] = []
    i = 0
    n = len(segment)
    while i < n:
        matched = False
        for surface, reading in pairs:
            length = len(surface)
            if length <= 0 or i + length > n:
                continue
            if segment[i : i + length] == surface:
                if fullwidth:
                    out.append("｛" + surface + "｜" + reading + "｝")
                else:
                    out.append("{" + surface + "|" + reading + "}")
                i += length
                matched = True
                break
        if not matched:
            out.append(segment[i])
            i += 1
    return "".join(out)


def to_fullwidth_ruby_delimiters(text: str) -> str:
    """半角ルビ {表記|よみ} を全角 ｛表記｜よみ｝ にそろえる。"""
    text = canonicalize_voicevox_ruby_delimiters(text or "")
    if not text:
        return ""

    def _repl(match: re.Match) -> str:
        surface = match.group(1)
        reading = normalize_voicevox_reading(match.group(2))
        return "｛" + surface + "｜" + reading + "｝"

    return _RUBY_TAG_RE.sub(_repl, text)


# 組み込みの最低限辞書（ファイルが無いときの予備）
DEFAULT_RUBY_DICT: list[tuple[str, str]] = [
    ("心筋梗塞", "しんきんこうそく"),
    ("心不全", "しんふぜん"),
    ("心房細動", "しんぼうさいどう"),
    ("心室細動", "しんしつさいどう"),
    ("心静止", "しんせいし"),
    ("肺塞栓", "はいそくせん"),
    ("敗血症", "はいけつしょう"),
    ("呼吸不全", "こきゅうふぜん"),
    ("気管内挿管", "きかんないそうかん"),
    ("気管挿管", "きかんそうかん"),
    ("胸骨圧迫", "きょうこつあっぱく"),
    ("昇圧剤", "しょうあつざい"),
    ("降圧", "こうあつ"),
    ("輸液", "ゆえき"),
    ("造影剤", "ぞうえいざい"),
    ("心電図", "しんでんず"),
    ("動脈血", "どうみゃくけつ"),
    ("静脈血", "じょうみゃくけつ"),
    ("酸素飽和度", "さんそほうわど"),
    ("意識障害", "いしきしょうがい"),
    ("昏睡", "こんすい"),
    ("痙攣", "けいれん"),
    ("麻痺", "まひ"),
    ("梗塞", "こうそく"),
    ("出血", "しゅっけつ"),
    ("麻酔", "ますい"),
    ("開腹", "かいふく"),
    ("開胸", "かいきょう"),
    ("縫合", "ほうごう"),
    ("抜管", "ばっかん"),
    ("挿管", "そうかん"),
    ("透析", "とうせき"),
    ("血糖", "けっとう"),
    ("白血球", "はっけっきゅう"),
    ("赤血球", "せっけっきゅう"),
    ("血小板", "けっしょうばん"),
    ("凝固", "ぎょうこ"),
    ("抗凝固", "こうぎょうこ"),
    ("抗生剤", "こうせいざい"),
    ("抗菌薬", "こうきんやく"),
]


def parse_ruby_dict_text(raw: str) -> list[tuple[str, str]]:
    """
    辞書テキストを読む。
    対応:
    - 用語<TAB>よみ
    - 用語,よみ
    - 用語｜よみ
    - 用語|よみ
    - 用語  よみ（空白2つ以上）
    - 用語 よみ（空白1つ。右側がかな中心のとき）
    """
    out: list[tuple[str, str]] = []
    seen: set[str] = set()
    for line in (raw or "").replace("\r\n", "\n").split("\n"):
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        surface = ""
        reading = ""
        if "\t" in line:
            surface, _, reading = line.partition("\t")
        elif "｜" in line:
            surface, _, reading = line.partition("｜")
        elif "|" in line and not line.startswith("{"):
            surface, _, reading = line.partition("|")
        elif "," in line:
            surface, _, reading = line.partition(",")
        else:
            parts = re.split(r"\s{2,}", line, maxsplit=1)
            if len(parts) == 2:
                surface, reading = parts
            else:
                # 1つの空白区切りも、右側が「よみ」らしいときは受け入れる
                parts = line.split()
                if len(parts) == 2:
                    maybe_surface, maybe_reading = parts
                    if re.fullmatch(r"[ぁ-んァ-ンー・ヴ゛゜A-Za-z0-9]+", maybe_reading):
                        surface, reading = maybe_surface, maybe_reading
        surface = surface.strip().strip("「」『』\"'")
        reading = normalize_voicevox_reading(reading)
        if not surface or not reading or surface in seen or surface == reading:
            continue
        seen.add(surface)
        out.append((surface, reading))
    return out


def load_ruby_dict_from_path(path: Path | None = None) -> list[tuple[str, str]]:
    """辞書ファイル（TSVなど）を読み込む。無ければ空リスト。"""
    p = path or RUBY_DICT_PATH
    try:
        if not p.is_file():
            return []
        return parse_ruby_dict_text(p.read_text(encoding="utf-8"))
    except OSError:
        return []


def get_active_ruby_dictionary() -> list[tuple[str, str]]:
    """
    使うルビ辞書を返す。
    優先: 編集で学習した語 → 画面で読み込んだ辞書 → 標準辞書 → 組み込み
    同じ用語は先勝ち。
    """
    merged: list[tuple[str, str]] = []
    seen: set[str] = set()

    def _add(pairs: list[tuple[str, str]]) -> None:
        for surface, reading in pairs:
            surface = (surface or "").strip()
            reading = normalize_voicevox_reading(reading)
            if not surface or not reading or surface in seen or surface == reading:
                continue
            seen.add(surface)
            merged.append((surface, reading))

    learned = None
    custom = None
    try:
        learned = st.session_state.get("ruby_dict_learned")
        custom = st.session_state.get("ruby_dict_custom")
    except Exception:
        learned = None
        custom = None
    if isinstance(learned, list) and learned:
        _add([(str(a), str(b)) for a, b in learned])
    if isinstance(custom, list) and custom:
        _add([(str(a), str(b)) for a, b in custom])
    _add(load_ruby_dict_from_path(RUBY_DICT_PATH))
    _add(DEFAULT_RUBY_DICT)
    return merged


def extract_ruby_pairs_from_script(script: str) -> list[tuple[str, str]]:
    """台本中の ｛用語｜よみ｝ をすべて取り出す（同じ用語は後勝ち）。"""
    text = canonicalize_voicevox_ruby_delimiters(script or "")
    order: list[str] = []
    mapping: dict[str, str] = {}
    for m in _RUBY_TAG_RE.finditer(text):
        surface = (m.group(1) or "").strip()
        reading = normalize_voicevox_reading(m.group(2) or "")
        if not surface or not reading or surface == reading:
            continue
        if surface not in mapping:
            order.append(surface)
        mapping[surface] = reading
    return [(s, mapping[s]) for s in order]


def is_ruby_dict_surface_allowed(surface: str) -> bool:
    """ルビ辞書へ学習反映してよい用語かを返す。1文字語は取り込まない。"""
    return len((surface or "").strip()) >= 2


def find_ruby_dict_updates(
    baseline_script: str,
    edited_script: str,
) -> list[tuple[str, str]]:
    """
    自動付与稿と修正稿を比べ、新規または読みが変わったルビを返す。
    """
    base_map = {s: r for s, r in extract_ruby_pairs_from_script(baseline_script)}
    updates: list[tuple[str, str]] = []
    for surface, reading in extract_ruby_pairs_from_script(edited_script):
        if not is_ruby_dict_surface_allowed(surface):
            continue
        if base_map.get(surface) != reading:
            updates.append((surface, reading))
    return updates


def find_rubies_missing_from_dictionary(
    script: str,
    dictionary: list[tuple[str, str]] | None = None,
) -> list[tuple[str, str]]:
    """
    ルビ入り原稿にあるが、辞書に無い（または読みが違う）ルビを返す。
    """
    dict_map = {
        s: r
        for s, r in (
            dictionary
            if dictionary is not None
            else get_active_ruby_dictionary()
        )
    }
    missing: list[tuple[str, str]] = []
    for surface, reading in extract_ruby_pairs_from_script(script):
        if not is_ruby_dict_surface_allowed(surface):
            continue
        if dict_map.get(surface) != reading:
            missing.append((surface, reading))
    return missing


def export_active_ruby_dictionary_file() -> Path:
    """いま有効なルビ辞書を ルビ辞書.txt に書き出す。"""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    export_text = format_ruby_dict_text(get_active_ruby_dictionary())
    RUBY_DICT_EXPORT_PATH.write_text(export_text, encoding="utf-8")
    st.session_state.ruby_dict_export_ready = True
    return RUBY_DICT_EXPORT_PATH


def sync_script_rubies_into_dictionary(script: str) -> list[tuple[str, str]]:
    """
    ルビ入り最終原稿と辞書を比べ、足りないルビを辞書へ追加する。
    追加後の辞書ファイルも書き出す。戻り値は今回追加した一覧。
    """
    missing = find_rubies_missing_from_dictionary(script)
    applied = apply_ruby_updates_to_learned_dict(missing) if missing else []
    # 追加がなくても、完成後にダウンロードできるよう最新版を書き出す
    export_active_ruby_dictionary_file()
    st.session_state.ruby_dict_post_video_updates = applied
    return applied


def format_ruby_dict_text(pairs: list[tuple[str, str]]) -> str:
    """用語\\tよみ 形式の辞書テキストにする。"""
    lines: list[str] = []
    seen: set[str] = set()
    for surface, reading in pairs:
        surface = (surface or "").strip()
        reading = normalize_voicevox_reading(reading)
        if not surface or not reading or surface == reading or surface in seen:
            continue
        seen.add(surface)
        lines.append(f"{surface}\t{reading}")
    return ("\n".join(lines) + "\n") if lines else ""


def apply_ruby_updates_to_learned_dict(
    updates: list[tuple[str, str]],
) -> list[tuple[str, str]]:
    """
    修正ルビを学習辞書に反映し、ルビ辞書.txt を書き出す。
    戻り値: 今回反映した更新一覧
    """
    if not updates:
        return []
    learned_raw = st.session_state.get("ruby_dict_learned") or []
    learned_map: dict[str, str] = {}
    order: list[str] = []
    for item in learned_raw:
        try:
            surface, reading = item[0], item[1]
        except Exception:
            continue
        surface = str(surface).strip()
        reading = normalize_voicevox_reading(str(reading))
        if not surface or not reading:
            continue
        if surface not in learned_map:
            order.append(surface)
        learned_map[surface] = reading
    applied: list[tuple[str, str]] = []
    for surface, reading in updates:
        surface = (surface or "").strip()
        reading = normalize_voicevox_reading(reading)
        if not surface or not reading or surface == reading:
            continue
        if not is_ruby_dict_surface_allowed(surface):
            continue
        if learned_map.get(surface) == reading:
            continue
        if surface not in learned_map:
            order.append(surface)
        learned_map[surface] = reading
        applied.append((surface, reading))
    st.session_state.ruby_dict_learned = [(s, learned_map[s]) for s in order]
    # いま有効な全辞書を書き出し（学習反映後）
    export_active_ruby_dictionary_file()
    st.session_state.ruby_dict_last_updates = applied
    return applied


def collect_dictionary_ruby_annotations(
    script: str,
    dictionary: list[tuple[str, str]] | None = None,
) -> list[dict[str, str]]:
    """
    台本中に現れる用語だけ、辞書からルビ候補を集める。
    文字数の多い用語を先に並べる（適用時の優先と揃える）。
    """
    text = strip_voicevox_ruby(script or "")
    dict_pairs = dictionary if dictionary is not None else get_active_ruby_dictionary()
    found: list[dict[str, str]] = []
    # 長い用語を先に（例: 「所見」を「所」より優先）
    for surface, reading in sorted(
        dict_pairs, key=lambda x: (len(x[0]), x[0]), reverse=True
    ):
        if surface in text and surface != reading:
            found.append({"surface": surface, "reading": reading})
    return found


def apply_dictionary_ruby_to_script(
    script: str,
    dictionary: list[tuple[str, str]] | None = None,
) -> tuple[str, int, list[dict[str, str]]]:
    """
    辞書を対照して ｛用語｜よみ｝ を付与する。
    同じ位置では文字数が多い用語を優先（例: 所 < 所見）。
    戻り値: (ルビ付き台本, 付与件数, 使った注釈一覧)
    """
    annotations = collect_dictionary_ruby_annotations(script, dictionary)
    out = apply_voicevox_ruby(script, annotations, fullwidth=True)
    return out, count_voicevox_ruby(out), annotations


def collect_default_ruby_annotations(script: str) -> list[dict[str, str]]:
    """互換: 辞書ルビ候補を集める。"""
    return collect_dictionary_ruby_annotations(script)


def merge_ruby_annotations(
    *groups: list[dict[str, str]],
) -> list[dict[str, str]]:
    """複数のルビ一覧をまとめ、同じ表記は先勝ち。"""
    merged: list[dict[str, str]] = []
    seen: set[str] = set()
    for group in groups:
        for item in group or []:
            surface = str(item.get("surface") or "").strip()
            reading = normalize_voicevox_reading(str(item.get("reading") or ""))
            if not surface or not reading or surface in seen or surface == reading:
                continue
            seen.add(surface)
            merged.append({"surface": surface, "reading": reading})
    return merged


def count_voicevox_ruby(text: str) -> int:
    return len(_RUBY_TAG_RE.findall(canonicalize_voicevox_ruby_delimiters(text or "")))


def prepare_script_for_voicevox(
    script: str,
    extra_annotations: list[dict[str, str]] | None = None,
    enabled: bool = True,
) -> tuple[str, int]:
    """
    音声生成直前にルビを整える。
    enabled=True（既定）: 辞書ルビ＋追加注釈（レビュー等）を付与
    enabled=False: ルビを外した文を返す
    戻り値: (台本, ルビ件数)
    """
    text = canonicalize_voicevox_ruby_delimiters(script or "")
    if not enabled:
        plain = strip_voicevox_ruby(text)
        return plain, 0
    # 辞書を優先（誤読の少ないよみ）→ そのあとレビュー等の追加注釈
    dict_ann = collect_dictionary_ruby_annotations(text)
    annotations = merge_ruby_annotations(dict_ann, list(extra_annotations or []))
    out = apply_voicevox_ruby(text, annotations, fullwidth=True)
    # VOICEVOX処理のため内部は半角にそろえて件数カウント
    out_half = canonicalize_voicevox_ruby_delimiters(out)
    return out_half, count_voicevox_ruby(out_half)


# 字幕折り返し／分割：行頭に置かない文字（閉じの 」 など）
SUBTITLE_NO_LINE_START = frozenset("」』）)]］】》〉、。，．！？!?ー…‥")
# 改ページ・折り返しで優先する切れ目（句読点など）
SUBTITLE_BREAK_AFTER = frozenset("。、，．！？!?…‥；;")
# 英単語・単位・薬物名など、途中で切りたくない文字
_LATIN_TOKEN_CHARS = frozenset(
    "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789./%-"
)


def _ruby_incomplete(fragment: str) -> bool:
    """断片の末尾で {表記|よみ} が途中切れなら True（半角/全角どちらも）。"""
    last_open = max(fragment.rfind("{"), fragment.rfind("｛"))
    if last_open < 0:
        return False
    tail = fragment[last_open:]
    return ("}" not in tail) and ("｝" not in tail)


def _find_preferred_break(buf: str) -> int:
    """
    折り返し候補の切れ目を返す（その位置より前を1行にする）。
    優先: 句読点の直後 → 空白の直後。見つからなければ -1。
    """
    if not buf:
        return -1
    for i in range(len(buf) - 1, -1, -1):
        if buf[i] in SUBTITLE_BREAK_AFTER and i + 1 < len(buf):
            # 行が極端に短くならない範囲で句読点直後を優先
            if i + 1 >= max(4, len(buf) // 4):
                return i + 1
    for i in range(len(buf) - 1, -1, -1):
        if buf[i].isspace() and i + 1 < len(buf):
            if i + 1 >= max(4, len(buf) // 4):
                return i + 1
    return -1


def _latin_token_start(buf: str) -> int:
    """buf 末尾が英数字トークンの途中なら、その開始位置。そうでなければ -1。"""
    if not buf or buf[-1] not in _LATIN_TOKEN_CHARS:
        return -1
    i = len(buf) - 1
    while i >= 0 and buf[i] in _LATIN_TOKEN_CHARS:
        i -= 1
    start = i + 1
    return start if start < len(buf) else -1


def _safe_force_chunks(text: str, max_chars: int) -> list[str]:
    """
    文字数で切るが、原則は句読点直後。
    ルビ途中・」直前・英単語／単位の途中では切らない。
    """
    if len(text) <= max_chars:
        return [text]
    out: list[str] = []
    i = 0
    n = len(text)
    while i < n:
        end = min(i + max_chars, n)
        # ルビ途中なら閉じるまで延ばす
        while end < n and _ruby_incomplete(text[i:end]):
            end += 1
        # 閉じ括弧・句読点が次チャンク先頭に来ないよう、同じチャンクに含める
        while end < n and text[end] in SUBTITLE_NO_LINE_START:
            end += 1
        # 可能なら句読点直後まで戻して切る（単語途中切断を避ける）
        if end < n:
            window = text[i:end]
            pref = _find_preferred_break(window)
            if pref > 0:
                end = i + pref
            else:
                # 英単語・単位の途中なら、トークン先頭まで戻す
                tok = _latin_token_start(window)
                if tok > 0:
                    end = i + tok
        if end <= i:
            end = min(i + 1, n)
        out.append(text[i:end])
        i = end
    return out


def is_katakana_notation_complaint(item: dict[str, str]) -> bool:
    """カタカナ表記そのものを問題にしている指摘なら True（除外用）。"""
    blob = " ".join(
        [
            str(item.get("issue") or ""),
            str(item.get("suggestion") or ""),
            str(item.get("original") or ""),
        ]
    )
    patterns = [
        r"カタカナ",
        r"漢字(に|へ|で|表記|に直|に直し|に修正|で書)",
        r"漢字表記",
        r"かな表記",
        r"仮名表記",
        r"読み仮名",
        r"ルビ",
        r"正式な漢字",
        r"漢字のほうが",
        r"漢字に(直し|変え|置換|修正)",
    ]
    return any(re.search(p, blob) for p in patterns)


def strip_code_fence(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json|JSON)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    return text.strip()


def repair_json_text(text: str) -> str:
    """よくある壊れ方を直してから json.loads する。"""
    text = strip_code_fence(text)
    # 最初の { から最後の } まで
    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        text = text[start : end + 1]
    # 末尾カンマ: ,] や ,} を除去
    text = re.sub(r",\s*([\]}])", r"\1", text)
    # スマートクォートを半角に
    text = (
        text.replace("“", '"')
        .replace("”", '"')
        .replace("‘", "'")
        .replace("’", "'")
    )
    return text


def parse_json_loose(text: str) -> dict[str, Any]:
    """モデル出力から JSON をできるだけ取り出す。"""
    candidates = [text, repair_json_text(text)]
    errors: list[str] = []
    for cand in candidates:
        try:
            data = json.loads(cand)
            if isinstance(data, dict):
                return data
        except json.JSONDecodeError as e:
            errors.append(str(e))

    # さらに: 制御文字除去して再試行
    cleaned = repair_json_text(re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text))
    try:
        data = json.loads(cleaned)
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError as e:
        errors.append(str(e))

    raise ValueError(
        "AIの返答をJSONとして読めませんでした。\n"
        + (errors[-1] if errors else "")
    )


def ask_claude_fix_json(api_key: str, model: str, broken: str) -> str:
    """壊れたJSONを、同じモデルに直してもらう。"""
    url = "https://api.anthropic.com/v1/messages"
    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    body = {
        "model": model,
        "max_tokens": 4096,
        "messages": [
            {
                "role": "user",
                "content": (
                    "次のテキストを、有効なJSONオブジェクトだけに直してください。"
                    "説明文やコードフェンスは不要です。"
                    "キーは medical_contradictions / awkward_for_doctors / "
                    "immersion_improvements / ruby_annotations を維持してください。"
                    "医学用語のカタカナ表記を問題にする指摘があれば削除してください。\n\n"
                    f"{broken[:12000]}"
                ),
            }
        ],
    }
    resp = http_session_direct().post(url, headers=headers, json=body, timeout=120)
    if resp.status_code != 200:
        raise RuntimeError(
            f"JSON修正リクエスト失敗 (HTTP {resp.status_code}): {resp.text[:300]}"
        )
    data = resp.json()
    parts = data.get("content", [])
    return "".join(p.get("text", "") for p in parts if p.get("type") == "text")


def http_session_direct() -> requests.Session:
    """
    プロキシ（通信の仲介役）を使わず、インターネットへ直接つなぐ。
    Cursor や社内ネットのプロキシ設定があると Claude API が
    403 Forbidden で失敗することがあるため。
    """
    session = requests.Session()
    session.trust_env = False  # HTTP_PROXY 等の環境変数を無視
    session.proxies = {"http": None, "https": None}
    return session


def review_with_claude(script: str, api_key: str) -> dict[str, Any]:
    """Anthropic Messages API でレビュー（APIキーは引数経由、直書きしない）。"""
    url = "https://api.anthropic.com/v1/messages"
    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    prompt = build_review_prompt(script)
    last_error = ""

    for model in CLAUDE_MODEL_CANDIDATES:
        body = {
            "model": model,
            "max_tokens": 8192,
            "messages": [{"role": "user", "content": prompt}],
        }
        try:
            resp = http_session_direct().post(
                url, headers=headers, json=body, timeout=120
            )
        except requests.exceptions.ProxyError as e:
            raise RuntimeError(
                "プロキシ（通信の仲介）のせいで Claude に接続できませんでした。\n"
                "ターミナルで次を実行してから、アプリを再起動してください:\n"
                "  unset HTTP_PROXY HTTPS_PROXY ALL_PROXY http_proxy https_proxy all_proxy\n"
                f"詳細: {e}"
            ) from e
        except requests.exceptions.RequestException as e:
            raise RuntimeError(
                "Claude API への通信に失敗しました。"
                "ネット接続と APIキーを確認してください。\n"
                f"詳細: {e}"
            ) from e

        if resp.status_code == 404 and "model" in resp.text.lower():
            last_error = f"{model}: {resp.text[:200]}"
            continue  # 次のモデル候補を試す

        if resp.status_code != 200:
            raise RuntimeError(
                f"Claude API エラー (HTTP {resp.status_code}): {resp.text[:500]}"
            )

        data = resp.json()
        parts = data.get("content", [])
        text = "".join(p.get("text", "") for p in parts if p.get("type") == "text")
        if not text:
            raise RuntimeError("Claude API から空の返答が返りました。")

        try:
            parsed = parse_json_loose(text)
        except ValueError:
            # 壊れたJSONを一度だけ修正依頼
            try:
                fixed = ask_claude_fix_json(api_key, model, text)
                parsed = parse_json_loose(fixed)
            except Exception as e:  # noqa: BLE001
                raise RuntimeError(
                    "Claudeのレビュー結果を読み取れませんでした。"
                    "もう一度『1. 台本をレビューする』を押してください。\n"
                    f"詳細: {e}"
                ) from e

        result = normalize_review(parsed)
        result["mode"] = "claude"
        result["model_used"] = model
        return result

    raise RuntimeError(
        "利用できる Claude モデルが見つかりませんでした。\n"
        f"試したモデル: {', '.join(CLAUDE_MODEL_CANDIDATES)}\n"
        f"最後のエラー: {last_error}"
    )


def build_drama_script_prompt(paper_text: str) -> str:
    """医学論文テキストから YouTube 医学ドラマ台本を作る指示文。"""
    paper = (paper_text or "").strip()
    if len(paper) > PAPER_TEXT_MAX_CHARS:
        paper = (
            paper[:PAPER_TEXT_MAX_CHARS]
            + "\n\n…（以下、長さ制限のため省略）"
        )
    return f"""あなたは医学教育向け YouTube 動画の台本作家かつ臨床医です。
次の医学論文（または症例報告）の内容をもとに、ナレーション台本だけを書いてください。

【必須条件】
① VOICEVOX 1.0倍速で読み上げたとき約10〜12分に収まる YouTube 医学ドラマ台本にする（本文の文字数はおおよそ {DRAMA_SCRIPT_TARGET_CHARS_MIN}〜{DRAMA_SCRIPT_TARGET_CHARS_MAX} 字を目安。これより長くしない）。
② すべてナレーターが話し、ドラマが展開する。場面転換もナレーションで示す。
③ 読み上げる台本本文以外は一切書かない。タイトル見出し、サブタイトル、シーン番号、「注釈」「解説」「制作メモ」、私への説明、前置き、後書き、Markdown記法は禁止。
④ 教育目的。ドラマ前半では正しい診断名を決して明示しない（示唆・鑑別の提示は可。確定診断は後半）。
⑤ 難易度は、医師免許を持つ研修医（初期研修医）が理解できるレベルにする。医学用語は使ってよいが、専門医向けの過度に高度な議論・稀少な略語の羅列は避ける。必要なら短い言い換えや文脈で意味が追えるようにする。ただし台本本文で視聴者に呼びかけない。「研修医のみなさん」「みなさん」「皆さん」などへの呼びかける表現は禁止（難易度の目安と、台詞・ナレーションの相手は別）。
⑥ 検査値は、医学的な意味付けが変わらない範囲で異なる数字に置き換えてよい（フィクション化）。ただし単位は原文の表記をそのまま使う（例: mg/dL, mEq/L, g/dL, IU/L）。単位を日本語訳や別表記に変えない。
⑦ 薬物名（一般名）はアルファベット表記のまま書く（例: vancomycin, meropenem）。カタカナ化や漢字訳にしない。
⑧ YouTube 字幕を想定し、各まとまりは短すぎず長すぎない長さ（おおよそ1画面に収まる程度）にする。
⑨ 登場人物のセリフには必ずカギ括弧「」を付ける。
⑩ 必ずしも一文ごとに改行しない。1画面の字幕に収まる長さなら、複数文を改行せずにつなげてよい。
⑪ 次の画面（改ページ／改行）へ移るときは、原則として句読点（。、！？ など）の直後で切る。単語の途中（漢字熟語の途中、英単語や単位・薬物名の途中）では切らない。
⑫ 鑑別診断の診断名を列挙するときは、診断名どうしを読点「、」だけで区切る。句点「。」では区切らない（禁止例: 心筋梗塞。心不全。肺炎。／正しい例: 心筋梗塞、心不全、肺炎。）。臓器名などの列挙も同様に読点でつなぎ、途中で改行しない。列挙全体の文末だけ句点で閉じる。例：胃、小腸、大腸、肝臓、胆嚢、骨盤内。
⑬ この段階ではルビ（｛用語｜よみ｝）を付けない。漢字のまま書く（ルビは後工程で辞書から付与する）。
⑭ 出力する前に、医師として医学的に違和感のある表現・論理の破綻を自分で直し、その最終稿だけを出す。

【出力形式】
- 台本本文のみ（プレーンテキスト）
- 最初の行からナレーションを始める
- コードブロックやJSONで囲まない

【論文テキスト】
{paper}
"""


def _claude_messages_text(
    api_key: str,
    prompt: str,
    *,
    max_tokens: int = 16000,
) -> tuple[str, str]:
    """Claude Messages API を呼び、返答テキストと使用モデル名を返す。"""
    url = "https://api.anthropic.com/v1/messages"
    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    last_error = ""
    for model in CLAUDE_MODEL_CANDIDATES:
        body = {
            "model": model,
            "max_tokens": int(max_tokens),
            "messages": [{"role": "user", "content": prompt}],
        }
        try:
            resp = http_session_direct().post(
                url, headers=headers, json=body, timeout=300
            )
        except requests.exceptions.ProxyError as e:
            raise RuntimeError(
                "プロキシのせいで Claude に接続できませんでした。\n"
                "ターミナルで proxy を解除してからアプリを再起動してください。\n"
                f"詳細: {e}"
            ) from e
        except requests.exceptions.RequestException as e:
            raise RuntimeError(
                "Claude API への通信に失敗しました。\n"
                f"詳細: {e}"
            ) from e

        if resp.status_code == 404 and "model" in resp.text.lower():
            last_error = f"{model}: {resp.text[:200]}"
            continue
        if resp.status_code != 200:
            raise RuntimeError(
                f"Claude API エラー (HTTP {resp.status_code}): {resp.text[:500]}"
            )
        data = resp.json()
        parts = data.get("content", [])
        text = "".join(
            p.get("text", "") for p in parts if p.get("type") == "text"
        ).strip()
        if not text:
            raise RuntimeError("Claude API から空の返答が返りました。")
        return text, model

    raise RuntimeError(
        "利用できる Claude モデルが見つかりませんでした。\n"
        f"試したモデル: {', '.join(CLAUDE_MODEL_CANDIDATES)}\n"
        f"最後のエラー: {last_error}"
    )


def _strip_script_wrappers(text: str) -> str:
    """台本以外の前置き・コード枠を取り除く。"""
    text = strip_code_fence((text or "").strip())
    # よくある前置き行を落とす
    lines = text.replace("\r\n", "\n").split("\n")
    drop_prefixes = (
        "以下が",
        "以下に",
        "台本を作成",
        "承知",
        "了解",
        "【台本】",
        "■台本",
        "# ",
    )
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines:
        first = lines[0].strip()
        if any(first.startswith(p) for p in drop_prefixes) and len(first) < 80:
            lines.pop(0)
            while lines and not lines[0].strip():
                lines.pop(0)
            continue
        break
    return "\n".join(lines).strip()


def polish_drama_script_medically(script: str, api_key: str) -> str:
    """医師視点で違和感・論理破綻を直し、台本本文だけ返す。"""
    prompt = f"""あなたは臨床医です。次の YouTube 医学ドラマ台本を読み、
医学的に違和感のある表現・論理的におかしい箇所だけを直してください。

【厳守】
- 出力は修正後の台本本文のみ（説明・箇条書きの修正リストは禁止）
- ナレーションのみの形式を維持する
- ドラマ前半で正しい診断を明示しないルールは維持する
- 難易度は医師免許を持つ研修医が理解できるレベルを維持する（過度に高度化しない）
- 「研修医のみなさん」「みなさん」など視聴者への呼びかけは入れない／残っていれば削除する
- 長さは VOICEVOX 1.0倍速で約10〜12分（おおよそ {DRAMA_SCRIPT_TARGET_CHARS_MIN}〜{DRAMA_SCRIPT_TARGET_CHARS_MAX} 字）を超えないよう、長くしすぎない
- セリフの「」、列挙の読点ルールは崩さない
- 鑑別診断の診断名を列挙するときは読点「、」で区切り、句点「。」では区切らない（文末の一句点だけ可）
- 検査値の単位は原文表記のまま（mg/dL, mEq/L など）。単位を書き換えない
- 薬物名はアルファベットのまま。カタカナ化しない
- 改行・改ページは原則として句読点の直後。単語・単位・薬物名の途中では切らない
- ルビ ｛用語｜よみ｝ は付けない／残っていれば外す（ルビは後工程）
- 不要な前置き・後書きを付けない

【台本】
{script}
"""
    text, _model = _claude_messages_text(api_key, prompt, max_tokens=16000)
    return _strip_script_wrappers(text)


def generate_drama_script_from_paper(paper_text: str, api_key: str) -> str:
    """
    医学論文テキストから、VOICEVOX 1倍速で約10〜12分のナレーション台本を作る。
    生成後に医学的な自己校正パスを1回行う。
    """
    paper = (paper_text or "").strip()
    if not paper:
        raise ValueError("論文テキストが空です。")
    if not (api_key or "").strip():
        raise RuntimeError(
            "論文から台本を作るには ANTHROPIC_API_KEY（Claude用の鍵）が必要です。"
            "画面上部でキーを入力・保存してください。"
        )
    draft, model_used = _claude_messages_text(
        api_key,
        build_drama_script_prompt(paper),
        max_tokens=16000,
    )
    draft = _strip_script_wrappers(draft)
    if not draft:
        raise RuntimeError("台本が空でした。もう一度お試しください。")
    try:
        polished = polish_drama_script_medically(draft, api_key)
        if polished.strip():
            draft = polished
    except Exception:
        # 校正に失敗しても下書きは返す
        pass
    _ = model_used
    return draft.strip()


def _heuristic_vancouver_from_paper(paper_text: str) -> str:
    """APIなし時の簡易抽出（DOIなどが見えるとき）。"""
    text = (paper_text or "").replace("\r\n", "\n")
    head = text[:8000]
    doi = ""
    m = re.search(
        r"(?:doi[:\s]*|https?://doi\.org/)(10\.\d{4,9}/[-._;()/:A-Z0-9]+)",
        head,
        flags=re.I,
    )
    if m:
        doi = m.group(1).rstrip(".")
    # 先頭付近のそれらしいタイトル行
    lines = [ln.strip() for ln in head.split("\n") if ln.strip()]
    title = ""
    for ln in lines[:40]:
        if len(ln) < 20 or len(ln) > 300:
            continue
        if re.search(r"abstract|introduction|keywords|©|copyright", ln, re.I):
            continue
        if re.search(r"[A-Za-z]{4,}", ln):
            title = ln
            break
    parts = []
    if title:
        parts.append(title.rstrip("."))
    if doi:
        parts.append(f"doi:{doi}")
    if parts:
        return ". ".join(parts) + ("." if not parts[-1].endswith(".") else "")
    return ""


def extract_vancouver_citation_from_paper(paper_text: str, api_key: str = "") -> str:
    """
    論文テキストから Vancouver 方式の参考文献1件を作る。
    例: Author A, Author B. Title. Journal. Year;Vol(Issue):Pages. doi:...
    """
    paper = (paper_text or "").strip()
    if not paper:
        return ""
    excerpt = paper[:PAPER_TEXT_MAX_CHARS]
    if (api_key or "").strip():
        prompt = f"""次の医学論文テキストから、エンディング画面用の参考文献を
Vancouver（バンクーバー）引用様式で1件だけ書いてください。

【形式の例】
Lim J, Wenham T. An Atypical Presentation of Mycoplasma pneumoniae Infection Mimicking Acute Surgical Abdomen in an Adult. Cureus. 2024;16(11):e73665. doi:10.7759/cureus.73665

【ルール】
- 出力は引用文1行（または必要なら2行）だけ。説明・箇条書き・前後の文言は禁止
- 著者は姓→名頭文字。3名超なら et al. を使ってよい
- 雑誌名・年・巻号・ページ／論文番号・DOI が分かれば入れる
- ライセンス（CC BY など）が本文にあれば末尾に「 / CC BY 4.0」のように付けてよい
- 不明な項目は無理に作らず省略する
- Markdownやコードブロックで囲まない

【論文テキスト】
{excerpt}
"""
        try:
            text, _ = _claude_messages_text(api_key, prompt, max_tokens=800)
            cite = _strip_script_wrappers(text)
            # 1〜3行に収める
            cite_lines = [ln.strip() for ln in cite.splitlines() if ln.strip()]
            cite = " ".join(cite_lines[:3]).strip()
            if cite:
                return cite
        except Exception:
            pass
    return _heuristic_vancouver_from_paper(paper)


def apply_paper_reference_to_session(citation: str) -> None:
    """参考文献をセッションとエンディング文面の初期値へ反映する。"""
    cite = (citation or "").strip()
    if not cite:
        return
    st.session_state.reference_text = cite
    try:
        save_reference_text(cite)
    except Exception:
        pass
    # エンディング④を最新の参考文献で入れ直せるようリセット
    st.session_state.ending_credits_text = ""
    st.session_state._ending_auto_text = ""
    st.session_state._ending_prefill_sig = None


def normalize_title_mukougawa(raw: str) -> str:
    """タイトルを『〜の向こう側』形に整える（すでに付いていればそのまま）。"""
    t = (raw or "").strip()
    t = t.strip("「」『』\"'").strip()
    t = re.sub(r"^(タイトル案|タイトル)[:：\s]*", "", t).strip()
    if "\n" in t:
        t = t.split("\n")[0].strip()
    t = re.sub(r"[。．!！?？]+$", "", t).strip()
    if not t:
        return "診断の向こう側"
    if t.endswith("の向こう側"):
        return t
    return f"{t}の向こう側"


def _heuristic_title_from_paper(paper_text: str) -> str:
    """APIなし時の簡潔なタイトル案。"""
    head = (paper_text or "")[:5000]
    # よくある疾患・病態キーワード（短いもの優先で拾う）
    keywords = [
        "敗血症",
        "心筋梗塞",
        "肺塞栓",
        "大動脈解離",
        "髄膜炎",
        "脳梗塞",
        "消化管穿孔",
        "急性腹症",
        "糖尿病性ケトアシドーシス",
        "アナフィラキシー",
        "心タンポナーデ",
        "気胸",
        "肺炎",
        "虫垂炎",
        "胆石",
        "膵炎",
        "腎盂腎炎",
        "腸閉塞",
        "心筋炎",
        "心不全",
        "喘息",
        "結核",
        "HIV",
        "SLE",
        "白血病",
        "リンパ腫",
    ]
    for kw in keywords:
        if kw in head:
            return normalize_title_mukougawa(kw)
    # 英語病名の簡易対応
    eng = [
        (r"\bsepsis\b", "敗血症"),
        (r"\bpneumonia\b", "肺炎"),
        (r"mycoplasma", "マイコプラズマ"),
        (r"appendicitis", "虫垂炎"),
        (r"pulmonary embolism", "肺塞栓"),
        (r"myocardial infarction", "心筋梗塞"),
    ]
    for pat, jp in eng:
        if re.search(pat, head, re.I):
            return normalize_title_mukougawa(jp)
    return "診断の向こう側"


def suggest_drama_title_from_paper(paper_text: str, api_key: str = "") -> str:
    """
    論文内容から簡潔なタイトル案を1つ作る。
    必ず『〜の向こう側』の形にする。
    """
    paper = (paper_text or "").strip()
    if not paper:
        return "診断の向こう側"
    excerpt = paper[: min(20000, len(paper))]
    if (api_key or "").strip():
        prompt = f"""次の医学論文（症例報告など）の内容から、
YouTube医学ドラマ用の簡潔な日本語タイトルを1つだけ考えてください。

【必須形式】
「〇〇の向こう側」
（例: 敗血症の向こう側 / 急性腹症の向こう側 / 肺塞栓の向こう側）

【ルール】
- 出力はタイトル1行だけ。説明・箇条書き・引用符・前置きは禁止
- 〇〇は疾患・病態・臨床テーマを短い日本語で（長くしすぎない）
- 論文タイトルの直訳にしない。視聴者が惹かれる簡潔な言葉にする
- 正しい診断名をネタバレしすぎない範囲で、テーマが伝わる語を選ぶ
- Markdownやコードブロックで囲まない

【論文テキスト】
{excerpt}
"""
        try:
            text, _ = _claude_messages_text(api_key, prompt, max_tokens=200)
            title = normalize_title_mukougawa(_strip_script_wrappers(text))
            if title:
                return title
        except Exception:
            pass
    return _heuristic_title_from_paper(paper)


def render_video_title_input() -> None:
    """
    タイトル手入力欄のみ（自動提案なし）。
    一度入れた内容は、ユーザーが上書きするまで session に残る。
    """
    st.text_input(
        "動画のタイトル（自分で入力）",
        key="video_title",
        placeholder="ここにタイトルを入力",
    )


def update_export_progress(
    progress,
    pct_box,
    status,
    pct: int,
    message: str = "",
) -> None:
    """進捗バーと％数字を同時に更新する（画面中央付近で大きく表示）。"""
    n = max(0, min(100, int(pct)))
    msg = (message or "").strip()
    bar_text = f"{n}%"
    if msg:
        bar_text = f"{n}%  {msg}"
    # Streamlit 1.50+: バー上にも％を出す
    try:
        progress.progress(n, text=bar_text)
    except TypeError:
        progress.progress(n / 100.0 if n <= 100 else 1.0)
    pct_box.markdown(
        f'<div style="font-size:2.4rem;font-weight:700;line-height:1.2;'
        f'margin:0.4rem 0 0.2rem 0;color:#111;">進捗 {n}%</div>',
        unsafe_allow_html=True,
    )
    if msg:
        status.info(f"{n}% — {msg}")
    else:
        status.info(f"{n}%")
    st.session_state.export_progress_pct = n
    st.session_state.export_progress_msg = msg


def heuristic_review(script: str) -> dict[str, Any]:
    """
    APIキーが無いとき用の簡易レビュー。
    本格的な医学監修の代わりではなく、プロセス確認用です。
    """
    medical: list[dict[str, str]] = []
    awkward: list[dict[str, str]] = []
    immersion: list[dict[str, str]] = []

    patterns_medical = [
        (
            r"血圧が\s*200[/\／]20",
            "血圧の下の値（拡張期）が極端に低すぎる表記は非現実的です。",
            "例: 「血圧 200/110」など臨床で起こりうる数値に修正してください。",
        ),
        (
            r"心電図(が|で)?(止まっ|フラット|直線)",
            "心停止の表現が曖昧です。波形の種類を明示した方が正確です。",
            "「心電図は心静止（アスystole）」「心室細動（VF）」など具体名に。",
        ),
        (
            r"酸素飽和度\s*(が)?\s*0%|SpO2\s*(が)?\s*0",
            "SpO2 0% は測定不能やプローブ外れの可能性が高く、物語上も説明が必要です。",
            "「測定不能」「プローブ外れの可能性」など状況説明を添えてください。",
        ),
    ]
    for pat, issue, suggestion in patterns_medical:
        m = re.search(pat, script)
        if m:
            medical.append(
                {
                    "original": m.group(0),
                    "issue": issue,
                    "suggestion": suggestion,
                }
            )

    patterns_awkward = [
        (
            r"オペ(を)?しよう|オペる",
            "現場では「オペ」単体より手技名・適応を言うことが多いです。",
            "「緊急開腹術を開始する」「気管内挿管する」など具体的に。",
        ),
        (
            r"点滴(を)?打(つ|って)",
            "医療者は「点滴を入れる／開始する」と言うことが多いです。",
            "「末梢ルートを確保して補液を開始」などに。",
        ),
        (
            r"心臓マッサージ",
            "現在は「胸骨圧迫」が標準的な言い方です。",
            "「胸骨圧迫を開始」に言い換えると医師の耳に自然です。",
        ),
    ]
    for pat, issue, suggestion in patterns_awkward:
        m = re.search(pat, script)
        if m:
            awkward.append(
                {
                    "original": m.group(0),
                    "issue": issue,
                    "suggestion": suggestion,
                }
            )

    if len(script) < 200:
        immersion.append(
            {
                "original": script[:80] + ("…" if len(script) > 80 else ""),
                "issue": "短いため、現場の音・時間経過・バイタルの変化が弱い可能性があります。",
                "suggestion": "モニター音、時刻、SpO2/血圧の推移、スタッフの短い掛け声を1〜2文足す。",
            }
        )
    if "…" not in script and "……" not in script and "——" not in script:
        immersion.append(
            {
                "original": "（全体）",
                "issue": "間（ま）や沈黙の描写が少なく、緊張感が平坦になりがちです。",
                "suggestion": "重要な決断の直前に短い沈黙やモニター音だけの一瞬を入れてください。",
            }
        )
    if not re.search(r"(血圧|SpO2|心拍数|脈拍|呼吸数)", script):
        immersion.append(
            {
                "original": "（バイタル表記なし）",
                "issue": "数値がないと救急シーンの臨場感が落ちます。",
                "suggestion": "「血圧 82/40、脈 130、SpO2 88%」など具体値を1か所入れてください。",
            }
        )

    if not medical and not awkward and not immersion:
        immersion.append(
            {
                "original": "（全体）",
                "issue": "自動チェックでは大きな問題は見つかりませんでした（簡易モード）。",
                "suggestion": "APIキーを設定すると、Claudeによる本格レビューに切り替えられます。",
            }
        )

    return normalize_review(
        {
            "medical_contradictions": medical,
            "awkward_for_doctors": awkward,
            "immersion_improvements": immersion,
            # 辞書ルビの自動付与はOFF
            "ruby_annotations": [],
            "mode": "heuristic",
        }
    )


def normalize_review(data: dict[str, Any]) -> dict[str, Any]:
    def _items(key: str) -> list[dict[str, str]]:
        raw = data.get(key, []) or []
        out: list[dict[str, str]] = []
        for item in raw:
            if not isinstance(item, dict):
                continue
            original = str(item.get("original", "")).strip()
            raw_suggestion = str(item.get("suggestion", "")).strip()
            cleaned = clean_script_replacement_text(raw_suggestion, original)
            normalized = {
                "original": original,
                "issue": str(item.get("issue", "")).strip(),
                # 解説付き修正案は本文だけ残す（反映ミス防止）
                "suggestion": cleaned or raw_suggestion,
                "suggestion_raw": raw_suggestion,
            }
            # カタカナ医学用語の表記指摘は採用しない
            if is_katakana_notation_complaint(normalized):
                continue
            out.append(normalized)
        return out

    ruby_out: list[dict[str, str]] = []
    seen_surface: set[str] = set()
    for item in data.get("ruby_annotations", []) or []:
        if not isinstance(item, dict):
            continue
        surface = str(item.get("surface", "")).strip()
        reading = normalize_voicevox_reading(str(item.get("reading", "")))
        if not surface or not reading or surface in seen_surface:
            continue
        seen_surface.add(surface)
        ruby_out.append({"surface": surface, "reading": reading})

    return {
        "medical_contradictions": _items("medical_contradictions"),
        "awkward_for_doctors": _items("awkward_for_doctors"),
        "immersion_improvements": _items("immersion_improvements"),
        "ruby_annotations": ruby_out,
        "mode": data.get("mode", "claude"),
    }


def run_script_review(script: str) -> dict[str, Any]:
    api_key = get_api_key()
    review_text = script
    truncated = False
    if len(script) > REVIEW_SCRIPT_MAX_CHARS:
        review_text = (
            script[:REVIEW_SCRIPT_MAX_CHARS]
            + "\n\n…（以下省略。レビューは先頭部分のみ）"
        )
        truncated = True
    if api_key:
        result = review_with_claude(review_text, api_key)
    else:
        result = heuristic_review(review_text)
    result["review_truncated"] = truncated

    # 辞書ルビを優先し、そのあと AI 提案ルビ（同じ用語は辞書が勝つ）
    dict_ann = collect_dictionary_ruby_annotations(script)
    ai_ann: list[dict[str, str]] = []
    seen: set[str] = set()
    for item in result.get("ruby_annotations") or []:
        surface = str(item.get("surface") or "").strip()
        if not surface or surface in seen:
            continue
        seen.add(surface)
        ai_ann.append(
            {
                "surface": surface,
                "reading": normalize_voicevox_reading(str(item.get("reading") or "")),
            }
        )
    merged = merge_ruby_annotations(dict_ann, ai_ann)
    result["ruby_annotations"] = merged
    result["script_with_ruby"] = apply_voicevox_ruby(script, merged, fullwidth=True)
    return result


# ---------------------------------------------------------------------------
# VOICEVOX
# ---------------------------------------------------------------------------
def check_voicevox() -> tuple[bool, str]:
    try:
        r = requests.get(f"{VOICEVOX_URL}/version", timeout=3)
        if r.status_code == 200:
            return True, r.text.strip().strip('"')
        return False, f"HTTP {r.status_code}"
    except requests.RequestException as e:
        return False, str(e)


def fetch_voicevox_speakers() -> list[dict[str, Any]]:
    """
    VOICEVOXにインストール済み（ダウンロード済み）の声優一覧を取得。
    各要素: {name, speaker_uuid, styles: [{name, id, type}, ...]}
    """
    try:
        r = requests.get(f"{VOICEVOX_URL}/speakers", timeout=8)
    except requests.RequestException as e:
        raise RuntimeError(f"VOICEVOXの声優一覧を取得できません: {e}") from e
    if r.status_code != 200:
        raise RuntimeError(
            f"VOICEVOXの声優一覧取得に失敗: HTTP {r.status_code} / {r.text[:200]}"
        )
    data = r.json()
    if not isinstance(data, list) or not data:
        raise RuntimeError("VOICEVOXに利用可能な声優がありません。")
    return data


def talk_styles_for_speaker(speaker: dict[str, Any]) -> list[dict[str, Any]]:
    """読み上げ用スタイルだけ（歌唱などは除外）。"""
    styles = speaker.get("styles") or []
    talk = [
        s
        for s in styles
        if isinstance(s, dict)
        and s.get("id") is not None
        and str(s.get("type") or "talk") == "talk"
    ]
    if talk:
        return talk
    # type が無い古いエンジン向け
    return [s for s in styles if isinstance(s, dict) and s.get("id") is not None]


def format_voicevox_credit(speaker_name: str, style_name: str = "") -> str:
    """エンディング用。声調名は表示せず、選択した声優名だけを出す。"""
    _ = style_name
    name = (speaker_name or "").strip() or DEFAULT_SPEAKER_NAME
    return f"VOICEVOX：{name}"


def build_ending_credits_text(reference_text: str, speaker_name: str) -> str:
    """
    エンディング全文を組み立てる（プレビュー／保存用）。
    参考文献・音声の直前は空行。区切りは文字の「----」ではなく描画時に線を引く。
    """
    ref = (reference_text or "").strip()
    if not ref:
        ref = "（参考文献未入力）"
    voice_line = format_voicevox_credit(speaker_name)
    return "\n".join(
        [
            ENDING_FICTION_NOTICE,
            "",
            "参考文献",
            "",
            ref,
            "",
            "音声",
            "",
            voice_line,
            "",
            ENDING_FOOTER,
        ]
    )


def create_ending_credits_frame(
    path: Path,
    ending_text: str = "",
    voicevox_credit: str = CREDIT_TEXT,
    reference_text: str | None = None,
    speaker_name: str | None = None,
) -> Path:
    """
    エンディング約10秒用。中央寄せ。
    参考文献・音声の前に余白を取り、区切りは横線で描いて文字と重ねない。
    """
    _ = voicevox_credit
    w, h = VIDEO_SIZE
    img = _dark_gradient_base((20, 25, 40))
    draw = ImageDraw.Draw(img)

    # 構造化データがあればそれを優先（重なり防止のため専用レイアウト）
    if reference_text is not None or speaker_name is not None:
        ref = (reference_text or "").strip() or "（参考文献未入力）"
        name = (speaker_name or "").strip() or DEFAULT_SPEAKER_NAME
        fiction_lines = [
            ln for ln in ENDING_FICTION_NOTICE.split("\n") if ln.strip()
        ]
        blocks: list[tuple[str, str]] = [
            ("text", fiction_lines[0] if fiction_lines else ""),
        ]
        for ln in fiction_lines[1:]:
            blocks.append(("text", ln))
        blocks.extend(
            [
                ("gap", ""),
                ("rule", ""),
                ("gap", ""),
                ("heading", "参考文献"),
                ("gap_sm", ""),
                ("text", ref),
                ("gap", ""),
                ("rule", ""),
                ("gap", ""),
                ("heading", "音声"),
                ("gap_sm", ""),
                ("text", format_voicevox_credit(name)),
                ("gap", ""),
                ("text", ENDING_FOOTER),
            ]
        )
    else:
        # 互換: 旧・全文テキスト（---- は無視して余白扱いにする）
        body = (ending_text or "").strip() or "（エンディング文未設定）"
        blocks = []
        for paragraph in body.replace("\r\n", "\n").split("\n"):
            p = paragraph.strip()
            if not p:
                blocks.append(("gap_sm", ""))
            elif set(p) <= set("-─━—–_") and len(p) >= 3:
                blocks.append(("gap", ""))
                blocks.append(("rule", ""))
                blocks.append(("gap", ""))
            elif p in ("参考文献", "音声"):
                blocks.append(("gap", ""))
                blocks.append(("heading", p))
                blocks.append(("gap_sm", ""))
            else:
                blocks.append(("text", p))

    # 全部同じ・小さめのフォント（見出しを大きくしない）
    body_font = load_jp_font(28, bold=False)
    max_w = int(w * 0.86)
    y = int(h * 0.08)
    fill = (220, 225, 230)

    def _draw_centered(line: str, font: ImageFont.ImageFont, color: tuple[int, int, int]) -> int:
        nonlocal y
        bbox = draw.textbbox((0, 0), line, font=font)
        tw = bbox[2] - bbox[0]
        th = max(bbox[3] - bbox[1], 1)
        x = (w - tw) // 2
        draw.text((x + 2, y + 2), line, font=font, fill=(0, 0, 0))
        draw.text((x, y), line, font=font, fill=color)
        y += th + 12
        return th

    for kind, content in blocks:
        if y > h - 48:
            break
        if kind == "gap":
            y += 40
            continue
        if kind == "gap_sm":
            y += 18
            continue
        if kind == "rule":
            # 横線の上下に余白を確保（見出し・本文と絶対に重ねない）
            y += 14
            rule_w = int(w * 0.22)
            x0 = (w - rule_w) // 2
            draw.line([(x0, y), (x0 + rule_w, y)], fill=(140, 150, 165), width=2)
            y += 36
            continue
        if kind == "heading":
            # 見出しも本文と同じ大きさ・同じ色
            _draw_centered(content, body_font, fill)
            continue
        # text: 長い行は折り返し、各行を中央寄せ
        wrapped = wrap_text_to_width(content, body_font, max_w, draw) or [""]
        # 英語の極端な1文字折り返しを減らす（スペース優先の再折り返し）
        if content.isascii() and " " in content:
            wrapped = _wrap_latin_prefer_spaces(content, body_font, max_w, draw)
        for line in wrapped:
            if y > h - 48:
                break
            if not line:
                y += 10
                continue
            _draw_centered(line, body_font, fill)

    path.parent.mkdir(parents=True, exist_ok=True)
    img.convert("RGB").save(path, format="PNG")
    return path


def resolve_default_voice_selection(
    speakers: list[dict[str, Any]],
) -> tuple[str, str, int]:
    """
    初期選択: できれば 青山龍星 ＋ ノーマル。無ければ先頭の声優＋ノーマル優先。
    戻り値: (声優名, 声調名, style_id)
    """
    by_name = {str(s.get("name") or ""): s for s in speakers}
    speaker = by_name.get(DEFAULT_SPEAKER_NAME) or speakers[0]
    speaker_name = str(speaker.get("name") or DEFAULT_SPEAKER_NAME)
    styles = talk_styles_for_speaker(speaker)
    if not styles:
        raise RuntimeError(f"「{speaker_name}」に読み上げ用の声調がありません。")
    preferred = (DEFAULT_STYLE_NAME, "ノーマル")
    style = None
    for want in preferred:
        for s in styles:
            if str(s.get("name") or "") == want:
                style = s
                break
        if style is not None:
            break
    if style is None:
        style = styles[0]
    return speaker_name, str(style.get("name") or "ノーマル"), int(style["id"])


def _split_after_punct_outside_quotes(text: str, punct: str) -> list[str]:
    """
    punct の直後で分割する。ただし「」『』のセリフ内では分割しない。
    （括弧内の「。」で切ると、閉じの 」 の前で字幕が改ページしてしまうため）
    """
    if not text:
        return []
    parts: list[str] = []
    buf: list[str] = []
    # いま開いているかぎ括弧の閉じ文字（入れ子対応）
    closers: list[str] = []
    for ch in text:
        buf.append(ch)
        if ch == "「":
            closers.append("」")
        elif ch == "『":
            closers.append("』")
        elif closers and ch == closers[-1]:
            closers.pop()
        elif (not closers) and (ch in punct):
            parts.append("".join(buf))
            buf = []
    if buf:
        parts.append("".join(buf))
    return parts


def _merge_chunks_broken_before_close(chunks: list[str]) -> list[str]:
    """閉じ括弧で始まる断片を直前の字幕に戻す（保険）。"""
    if not chunks:
        return []
    out: list[str] = [chunks[0]]
    for c in chunks[1:]:
        if c and c[0] in "」』" and out:
            out[-1] += c
        else:
            out.append(c)
    return out


def split_text_for_voicevox(text: str, max_chars: int = MAX_VOICEVOX_CHARS) -> list[str]:
    """
    句点単位で分割する（文章同士はくっつけない）。
    理由: 1音声区間＝1字幕にすると、読み上げと表示がズレにくい。
    「」『』のセリフ内では句点・読点で切らない（閉じ括弧の前で改ページしない）。
    1文が長すぎるときだけ max_chars でさらに切る（ルビの途中・」直前では切らない）。
    """
    text = text.replace("\r\n", "\n").strip()
    if not text:
        return []
    blocks = re.split(r"\n+", text)
    chunks: list[str] = []
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        # 「」内の 。！？ では切らない
        parts = _split_after_punct_outside_quotes(block, "。！？!?")
        for part in parts:
            part = part.strip()
            if not part:
                continue
            if len(part) <= max_chars:
                chunks.append(part)
            else:
                # 読点でも切れなければ文字数で強制分割（「」内の読点では切らない）
                buf = ""
                for piece in _split_after_punct_outside_quotes(part, "、,，"):
                    piece = piece.strip()
                    if not piece:
                        continue
                    if len(buf) + len(piece) <= max_chars:
                        buf += piece
                    else:
                        if buf:
                            chunks.append(buf)
                        if len(piece) <= max_chars:
                            buf = piece
                        else:
                            chunks.extend(_safe_force_chunks(piece, max_chars))
                            buf = ""
                if buf:
                    chunks.append(buf)
    return _merge_chunks_broken_before_close(chunks)


def clamp_voicevox_speed(speed: float) -> float:
    """読み上げ速度を 0.8〜1.5 の範囲に収める（0.1 刻み）。"""
    try:
        s = float(speed)
    except (TypeError, ValueError):
        s = float(VOICEVOX_SPEED_SCALE)
    s = max(VOICEVOX_SPEED_MIN, min(VOICEVOX_SPEED_MAX, s))
    return round(round(s / VOICEVOX_SPEED_STEP) * VOICEVOX_SPEED_STEP, 1)


def synthesize_wav_bytes(
    text: str,
    speaker: int = DEFAULT_SPEAKER_ID,
    speed_scale: float = VOICEVOX_SPEED_SCALE,
) -> bytes:
    q = requests.post(
        f"{VOICEVOX_URL}/audio_query",
        params={"text": text, "speaker": speaker},
        timeout=60,
    )
    if q.status_code != 200:
        raise RuntimeError(f"audio_query 失敗: HTTP {q.status_code} / {q.text[:300]}")
    query = q.json()
    speed = clamp_voicevox_speed(speed_scale)
    query["speedScale"] = float(query.get("speedScale", 1.0)) * speed
    query["intonationScale"] = float(query.get("intonationScale", 1.0))

    s = requests.post(
        f"{VOICEVOX_URL}/synthesis",
        params={"speaker": speaker},
        headers={"Content-Type": "application/json"},
        data=json.dumps(query),
        timeout=120,
    )
    if s.status_code != 200:
        raise RuntimeError(f"synthesis 失敗: HTTP {s.status_code} / {s.text[:300]}")
    return s.content


def concat_wav_files(wav_paths: list[Path], out_path: Path, pause_ms: int = 350) -> Path:
    """複数WAVを無音つきでつなぎ、ファイルへ書き出す（長尺向け・メモリ節約）。"""
    if not wav_paths:
        raise ValueError("音声データが空です。")

    params = None
    with wave.open(str(out_path), "wb") as out_w:
        for i, p in enumerate(wav_paths):
            with wave.open(str(p), "rb") as w:
                if params is None:
                    params = w.getparams()
                    out_w.setparams(params)
                else:
                    if (
                        w.getnchannels() != params.nchannels
                        or w.getsampwidth() != params.sampwidth
                        or w.getframerate() != params.framerate
                    ):
                        raise RuntimeError("WAVの形式が一致しません。")
                out_w.writeframes(w.readframes(w.getnframes()))
                if i < len(wav_paths) - 1 and pause_ms > 0 and params is not None:
                    n = int(params.framerate * pause_ms / 1000)
                    silence = b"\x00" * n * params.nchannels * params.sampwidth
                    out_w.writeframes(silence)
    return out_path


def prepend_silence_to_wav(wav_path: Path, silence_sec: float) -> Path:
    """既存WAVの先頭に無音を足す（朗読開始までの余白用）。"""
    silence_sec = max(0.0, float(silence_sec))
    if silence_sec <= 0:
        return wav_path
    tmp = wav_path.with_suffix(".intro_silence.tmp.wav")
    with wave.open(str(wav_path), "rb") as src:
        params = src.getparams()
        body = src.readframes(src.getnframes())
    n_silent = int(params.framerate * silence_sec)
    silence = b"\x00" * n_silent * params.nchannels * params.sampwidth
    with wave.open(str(tmp), "wb") as out_w:
        out_w.setparams(params)
        out_w.writeframes(silence)
        out_w.writeframes(body)
    tmp.replace(wav_path)
    return wav_path


def shift_subtitle_cues(
    cues: list[dict[str, Any]], offset_sec: float
) -> list[dict[str, Any]]:
    """字幕の開始・終了時刻をまとめてずらす（冒頭無音と同期させる）。"""
    offset_sec = float(offset_sec)
    if abs(offset_sec) < 1e-9:
        return cues
    for cue in cues:
        cue["start"] = float(cue.get("start", 0)) + offset_sec
        cue["end"] = float(cue.get("end", 0)) + offset_sec
    return cues


def strip_voicevox_ruby(text: str) -> str:
    """
    VOICEVOXルビ {表記|よみ} / ｛表記｜よみ｝ などを外し、
    字幕・表示用に「表記」だけ残す（半角/全角の区切りは区別しない）。
    """
    text = canonicalize_voicevox_ruby_delimiters(text or "")
    if not text:
        return ""
    return re.sub(
        r"\{([^|\n]+)\|[^}\n]+\}",
        r"\1",
        text,
    )


def expand_voicevox_ruby_to_reading(text: str) -> str:
    """
    VOICEVOXルビ {表記|よみ} / ｛表記｜よみ｝ などを外し、
    読み上げ用に「よみ」だけ残す（半角/全角の区切りは区別しない）。
    """
    text = canonicalize_voicevox_ruby_delimiters(text or "")
    if not text:
        return ""
    return re.sub(
        r"\{[^|\n]+\|([^}\n]+)\}",
        r"\1",
        text,
    )


def create_subtitle_png(text: str, path: Path) -> Path:
    """画面下部中央に載せる字幕PNG（透過・縁取り）。"""
    w, h = VIDEO_SIZE
    img = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)
    font = load_jp_font(52, bold=True)
    lines = wrap_text_to_width((text or "").strip(), font, int(w * 0.86), draw)[:3]
    if not lines:
        path.parent.mkdir(parents=True, exist_ok=True)
        img.save(path, format="PNG")
        return path

    line_sizes = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        line_sizes.append((bbox[2] - bbox[0], bbox[3] - bbox[1]))
    pad_x, pad_y = 28, 18
    box_w = max(s[0] for s in line_sizes) + pad_x * 2
    box_h = sum(s[1] for s in line_sizes) + 12 * (len(lines) - 1) + pad_y * 2
    box_x = (w - box_w) // 2
    # 注意書きより上
    box_y = int(h * 0.72) - box_h

    overlay = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    od = ImageDraw.Draw(overlay)
    try:
        od.rounded_rectangle(
            [box_x, box_y, box_x + box_w, box_y + box_h],
            radius=16,
            fill=(0, 0, 0, 160),
        )
    except Exception:
        od.rectangle(
            [box_x, box_y, box_x + box_w, box_y + box_h],
            fill=(0, 0, 0, 160),
        )
    img = Image.alpha_composite(img, overlay)
    draw = ImageDraw.Draw(img)

    y = box_y + pad_y
    for line, (lw, lh) in zip(lines, line_sizes):
        x = (w - lw) // 2
        draw_outlined_text(
            draw,
            (x, y),
            line,
            font,
            fill=(255, 255, 255),
            outline=(0, 0, 0),
            outline_width=4,
        )
        y += lh + 12

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, format="PNG")
    return path


def generate_narration_wav_to_file(
    script: str,
    out_wav: Path,
    progress_callback=None,
    pause_ms: int = 280,
    speaker: int = DEFAULT_SPEAKER_ID,
    speed_scale: float = VOICEVOX_SPEED_SCALE,
) -> tuple[Path, list[dict[str, Any]]]:
    """
    長い台本向け: VOICEVOXで音声生成し、字幕用タイミングも返す。
    1音声区間＝1字幕。開始・終了は各WAVの実時間（よみ上げと完全同期）。
    戻り値: (wavパス, 字幕キュー[{start,end,text}, ...])
    text は表記のみ（ルビ記号なし）。
    """
    chunks = split_text_for_voicevox(script)
    if not chunks:
        raise ValueError("読み上げる文章が空です。")

    out_wav.parent.mkdir(parents=True, exist_ok=True)
    part_dir = out_wav.parent / "_voice_parts"
    part_dir.mkdir(parents=True, exist_ok=True)
    part_paths: list[Path] = []
    subtitle_cues: list[dict[str, Any]] = []
    t = 0.0
    pause_sec = max(0.0, pause_ms / 1000.0)
    speed = clamp_voicevox_speed(speed_scale)

    try:
        for i, chunk in enumerate(chunks):
            if progress_callback:
                progress_callback(i, len(chunks))
            part = part_dir / f"part_{i:05d}.wav"
            # VOICEVOXへはよみがなのみ、字幕には表記のみ（同じ区間の実時間を共有）
            tts_text = expand_voicevox_ruby_to_reading(chunk).strip()
            display = strip_voicevox_ruby(chunk).strip()
            if not tts_text:
                tts_text = display
            if not tts_text:
                # 空区間はスキップ（無音も字幕も入れない）
                continue
            # 連結時と同じく、前の音声のあとにだけ無音を入れる
            if part_paths:
                t += pause_sec
            part.write_bytes(
                synthesize_wav_bytes(tts_text, speaker=speaker, speed_scale=speed)
            )
            part_paths.append(part)
            with wave.open(str(part), "rb") as w:
                dur = w.getnframes() / float(w.getframerate())
            seg_start = t
            seg_end = t + dur
            if display and dur > 0:
                subtitle_cues.append(
                    {
                        "start": seg_start,
                        "end": seg_end,
                        "text": display,
                        # 検証用: この字幕が対応する読み上げ文
                        "tts": tts_text,
                    }
                )
            t = seg_end
        if progress_callback:
            progress_callback(len(chunks), len(chunks))
        if not part_paths:
            raise ValueError("読み上げる文章が空です。")
        concat_wav_files(part_paths, out_wav, pause_ms=pause_ms)

        # 連結WAVの実時間とタイムライン t を一致させる（字幕も同じ倍率）
        with wave.open(str(out_wav), "rb") as w:
            actual_dur = w.getnframes() / float(w.getframerate())
        if subtitle_cues and t > 0 and abs(actual_dur - t) > 0.001:
            scale = actual_dur / t
            for cue in subtitle_cues:
                cue["start"] = float(cue["start"]) * scale
                cue["end"] = float(cue["end"]) * scale
            t = actual_dur
    finally:
        for p in part_paths:
            try:
                p.unlink(missing_ok=True)
            except OSError:
                pass
        try:
            part_dir.rmdir()
        except OSError:
            pass
    return out_wav, subtitle_cues


def validate_audio_subtitle_sync(
    wav_path: Path,
    subtitle_cues: list[dict[str, Any]],
    tol_sec: float = 0.08,
) -> list[str]:
    """
    音声WAVと字幕キューが同期しているか検査する。
    問題があればメッセージ一覧を返す（空ならOK）。
    """
    issues: list[str] = []
    if not wav_path.is_file():
        return ["音声ファイルがありません"]
    with wave.open(str(wav_path), "rb") as w:
        wav_dur = w.getnframes() / float(w.getframerate())
    if wav_dur <= 0:
        return ["音声の長さが 0 です"]
    if not subtitle_cues:
        issues.append("字幕キューが空です")
        return issues

    first_start = float(subtitle_cues[0].get("start", 0))
    if first_start > tol_sec:
        issues.append(
            f"最初の字幕開始が遅すぎます ({first_start:.3f}s)"
        )

    prev_end = 0.0
    for i, cue in enumerate(subtitle_cues):
        start = float(cue.get("start", 0))
        end = float(cue.get("end", 0))
        text = str(cue.get("text") or "").strip()
        if not text:
            issues.append(f"字幕#{i+1}: テキストが空です")
        if end <= start:
            issues.append(f"字幕#{i+1}: 終了が開始以下です ({start:.3f}→{end:.3f})")
        if start < -tol_sec:
            issues.append(f"字幕#{i+1}: 開始が負です ({start:.3f})")
        if end > wav_dur + tol_sec:
            issues.append(
                f"字幕#{i+1}: 音声長を超えています"
                f" (終了 {end:.3f}s / 音声 {wav_dur:.3f}s)"
            )
        if i > 0 and start < prev_end - tol_sec:
            issues.append(
                f"字幕#{i+1}: 前の字幕と重なっています"
                f" (前終了 {prev_end:.3f} / 開始 {start:.3f})"
            )
        prev_end = max(prev_end, end)

    last_end = float(subtitle_cues[-1].get("end", 0))
    # 末尾に字幕なしの読み上げが無いとき、最後の字幕終了は音声長に近い
    if abs(last_end - wav_dur) > tol_sec and last_end < wav_dur - 1.0:
        # 1秒以上余る場合のみ注意（末尾の表示なし区間が長い）
        issues.append(
            f"注意: 最後の字幕終了 ({last_end:.3f}s) のあと音声が"
            f" {wav_dur - last_end:.3f}s 続きます"
        )
    return issues


def generate_narration_wav(script: str) -> bytes:
    """短い用途向け互換。長尺は generate_narration_wav_to_file を使う。"""
    with tempfile.TemporaryDirectory(prefix="vvox_") as tmp:
        path = Path(tmp) / "n.wav"
        generate_narration_wav_to_file(script, path)
        return path.read_bytes()


# ---------------------------------------------------------------------------
# 背景画像（Pillow）— YouTubeサムネ風タイトル＋VOICEVOXクレジット
# ---------------------------------------------------------------------------
def load_jp_font(size: int, bold: bool = True) -> ImageFont.ImageFont:
    """macOS で使える日本語フォントを探す。"""
    candidates = []
    if bold:
        candidates.extend(
            [
                "/System/Library/Fonts/ヒラギノ角ゴシック W8.ttc",
                "/System/Library/Fonts/ヒラギノ角ゴシック W6.ttc",
                "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            ]
        )
    candidates.extend(
        [
            "/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc",
            "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            "/Library/Fonts/Arial Unicode.ttf",
        ]
    )
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_outlined_text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: tuple[int, int, int],
    outline: tuple[int, int, int] = (0, 0, 0),
    outline_width: int = 6,
) -> None:
    """縁取り＋影つきの太字テキスト（YouTubeサムネ風）。"""
    x, y = xy
    # 影
    draw.text((x + 4, y + 4), text, font=font, fill=(0, 0, 0))
    # 縁取り（周囲に同じ文字を重ねる）
    for dx in range(-outline_width, outline_width + 1):
        for dy in range(-outline_width, outline_width + 1):
            if dx == 0 and dy == 0:
                continue
            if dx * dx + dy * dy > outline_width * outline_width:
                continue
            draw.text((x + dx, y + dy), text, font=font, fill=outline)
    draw.text((x, y), text, font=font, fill=fill)


# 字幕折り返し：行頭に置かない文字は SUBTITLE_NO_LINE_START を参照


def wrap_text_to_width(
    text: str, font: ImageFont.ImageFont, max_width: int, draw: ImageDraw.ImageDraw
) -> list[str]:
    """
    画面幅に収まるよう折り返す。
    原則: 句読点の直後で改行。英単語・単位・薬物名の途中では切らない。
    「」で囲んだセリフでは、閉じの 」 の直前では改行しない（行頭禁則）。
    """
    text = (text or "").replace("\r\n", "\n").strip()
    if not text:
        return []
    lines: list[str] = []
    for paragraph in text.split("\n"):
        if not paragraph:
            lines.append("")
            continue
        buf = ""
        for ch in paragraph:
            trial = buf + ch
            bbox = draw.textbbox((0, 0), trial, font=font)
            if bbox[2] - bbox[0] <= max_width:
                buf = trial
                continue
            # 次の文字が 」 など行頭禁則 → 直前では切らず、同じ行に付ける
            if ch in SUBTITLE_NO_LINE_START and buf:
                lines.append(buf + ch)
                buf = ""
                continue
            # 句読点直後（なければ空白）で切る
            pref = _find_preferred_break(buf)
            if pref > 0:
                lines.append(buf[:pref])
                buf = buf[pref:] + ch
                continue
            # 英単語・単位の途中なら、トークンごと次行へ
            tok = _latin_token_start(buf)
            if tok > 0 and ch in _LATIN_TOKEN_CHARS:
                lines.append(buf[:tok])
                buf = buf[tok:] + ch
                continue
            if buf:
                lines.append(buf)
            buf = ch
        if buf:
            lines.append(buf)
    return lines


def _wrap_latin_prefer_spaces(
    text: str, font: ImageFont.ImageFont, max_width: int, draw: ImageDraw.ImageDraw
) -> list[str]:
    """英語など、できれば単語の途中ではなくスペースで折り返す。"""
    words = (text or "").split()
    if not words:
        return []
    lines: list[str] = []
    buf = ""
    for word in words:
        trial = word if not buf else f"{buf} {word}"
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= max_width:
            buf = trial
            continue
        if buf:
            lines.append(buf)
        # 単語自体が幅を超える場合は1文字折り返し
        if draw.textbbox((0, 0), word, font=font)[2] > max_width:
            lines.extend(wrap_text_to_width(word, font, max_width, draw))
            buf = ""
        else:
            buf = word
    if buf:
        lines.append(buf)
    return lines


def fit_image_cover(img: Image.Image, size: tuple[int, int]) -> Image.Image:
    """画像を 1920x1080 に合わせて中央切り抜き（余白なし）。"""
    tw, th = size
    iw, ih = img.size
    scale = max(tw / iw, th / ih)
    nw, nh = int(iw * scale), int(ih * scale)
    img = img.resize((nw, nh), Image.Resampling.LANCZOS)
    left = (nw - tw) // 2
    top = (nh - th) // 2
    return img.crop((left, top, left + tw, top + th))


def load_image_from_upload(uploaded_file) -> Image.Image | None:
    if uploaded_file is None:
        return None
    data = uploaded_file.getvalue() if hasattr(uploaded_file, "getvalue") else uploaded_file.read()
    return Image.open(io.BytesIO(data)).convert("RGBA")


def load_text_from_upload(uploaded_file) -> str:
    """アップロードファイルからテキストを読む（.txt / .docx 両対応）。"""
    if uploaded_file is None:
        return ""
    name = uploaded_file.name or "upload.txt"
    raw = (
        uploaded_file.getvalue()
        if hasattr(uploaded_file, "getvalue")
        else uploaded_file.read()
    )
    if isinstance(raw, str):
        raw = raw.encode("utf-8")
    lower = name.lower()
    if not (
        lower.endswith(".txt")
        or lower.endswith(".text")
        or lower.endswith(".md")
        or lower.endswith(".docx")
    ):
        # 拡張子が無い／不明なときは中身で判断
        name = "upload.docx" if raw[:2] == b"PK" else "upload.txt"
    return extract_text_from_bytes(name, raw)


def text_to_docx_bytes(text: str) -> bytes:
    """プレーンテキストを Word（.docx）のバイト列にする。"""
    doc = Document()
    lines = (text or "").replace("\r\n", "\n").split("\n")
    if not lines:
        doc.add_paragraph("")
    else:
        for line in lines:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def paste_centered(
    base: Image.Image,
    overlay: Image.Image,
    max_width_ratio: float = 0.9,
    max_height_ratio: float = 0.35,
    y_ratio: float = 0.08,
) -> None:
    """タイトル画像などを、幅制限つきで中央に重ねる。"""
    w, h = base.size
    max_w = int(w * max_width_ratio)
    max_h = int(h * max_height_ratio)
    ow, oh = overlay.size
    scale = min(max_w / ow, max_h / oh, 1.0)
    nw, nh = max(1, int(ow * scale)), max(1, int(oh * scale))
    overlay = overlay.resize((nw, nh), Image.Resampling.LANCZOS)
    x = (w - nw) // 2
    y = int(h * y_ratio)
    if overlay.mode != "RGBA":
        overlay = overlay.convert("RGBA")
    base.paste(overlay, (x, y), overlay)


def _dark_gradient_base(accent: tuple[int, int, int] = (20, 50, 80)) -> Image.Image:
    """著作権フリーの暗い医療背景ベース（自作図形のみ）。"""
    w, h = VIDEO_SIZE
    img = Image.new("RGB", (w, h), (6, 10, 18))
    draw = ImageDraw.Draw(img)
    ar, ag, ab = accent
    for y in range(h):
        t = y / max(h - 1, 1)
        r = int(6 + ar * 0.35 * t)
        g = int(10 + ag * 0.35 * t)
        b = int(18 + ab * 0.45 * t)
        draw.line([(0, y), (w, y)], fill=(r, g, b))
    return img


def _draw_ecg_wave(draw: ImageDraw.ImageDraw, y0: int, color=(80, 220, 160)) -> None:
    w, _ = VIDEO_SIZE
    points = []
    x = 80
    while x < w - 80:
        points.extend(
            [
                (x, y0),
                (x + 20, y0),
                (x + 28, y0 - 18),
                (x + 36, y0 + 70),
                (x + 44, y0 - 90),
                (x + 52, y0 + 20),
                (x + 60, y0),
                (x + 100, y0),
            ]
        )
        x += 120
    if len(points) >= 2:
        draw.line(points, fill=color, width=3)


def draw_medical_scene(theme: str) -> Image.Image:
    """シンプルな医学イメージを自作（写真不使用＝著作権フリー）。"""
    w, h = VIDEO_SIZE
    themes = {
        "er": ((40, 20, 30), "救急・初療"),
        "icu": ((15, 35, 55), "集中治療"),
        "surgery": ((25, 25, 40), "手術・処置"),
        "lab": ((20, 40, 45), "検査・画像"),
        "ward": ((25, 35, 30), "病棟・経過"),
        "consult": ((35, 30, 45), "説明・決断"),
        "pharma": ((30, 40, 35), "治療・投与"),
        "ambulance": ((45, 25, 25), "搬送・現場"),
    }
    accent, label = themes.get(theme, themes["icu"])
    img = _dark_gradient_base(accent)
    draw = ImageDraw.Draw(img)
    cx, cy = w // 2, int(h * 0.42)

    if theme == "er":
        draw.rectangle([cx - 420, cy - 220, cx + 420, cy + 220], outline=(180, 60, 70), width=4)
        draw.line([(cx - 200, cy), (cx + 200, cy)], fill=(200, 70, 80), width=6)
        draw.line([(cx, cy - 160), (cx, cy + 160)], fill=(200, 70, 80), width=6)
        _draw_ecg_wave(draw, cy + 40, (220, 120, 120))
    elif theme == "icu":
        draw.ellipse([cx - 260, cy - 260, cx + 260, cy + 260], outline=(50, 120, 160), width=3)
        draw.rectangle([cx - 380, cy - 200, cx + 380, cy + 200], outline=(70, 140, 180), width=3)
        _draw_ecg_wave(draw, cy, (80, 220, 180))
        for i, val in enumerate(["HR", "BP", "SpO2"]):
            x = cx - 300 + i * 220
            draw.text((x, cy + 120), val, fill=(160, 200, 220), font=load_jp_font(36, bold=True))
    elif theme == "surgery":
        draw.ellipse([cx - 180, cy - 120, cx + 180, cy + 120], outline=(200, 200, 210), width=5)
        draw.arc([cx - 280, cy - 200, cx + 280, cy + 200], 200, 340, fill=(160, 170, 190), width=8)
        draw.line([(cx - 40, cy - 40), (cx + 120, cy + 80)], fill=(210, 210, 220), width=5)
    elif theme == "lab":
        for i in range(3):
            x0 = cx - 420 + i * 300
            draw.rectangle([x0, cy - 180, x0 + 240, cy + 180], outline=(100, 180, 190), width=3)
            draw.line([(x0 + 20, cy - 100), (x0 + 220, cy + 100)], fill=(80, 150, 160), width=2)
            draw.line([(x0 + 20, cy + 80), (x0 + 220, cy - 60)], fill=(80, 150, 160), width=2)
    elif theme == "ward":
        try:
            draw.rounded_rectangle(
                [cx - 400, cy - 40, cx + 400, cy + 120],
                radius=30,
                outline=(90, 140, 110),
                width=4,
            )
        except Exception:
            draw.rectangle([cx - 400, cy - 40, cx + 400, cy + 120], outline=(90, 140, 110), width=4)
        draw.rectangle([cx - 380, cy - 160, cx - 300, cy - 40], outline=(90, 140, 110), width=3)
        _draw_ecg_wave(draw, cy - 100, (100, 190, 140))
    elif theme == "consult":
        draw.ellipse([cx - 320, cy - 80, cx - 120, cy + 120], outline=(160, 140, 200), width=4)
        draw.ellipse([cx + 120, cy - 80, cx + 320, cy + 120], outline=(140, 160, 210), width=4)
        draw.line([(cx - 100, cy + 20), (cx + 100, cy + 20)], fill=(170, 160, 200), width=3)
    elif theme == "pharma":
        draw.line([(cx - 80, cy - 220), (cx - 80, cy + 200)], fill=(140, 190, 160), width=6)
        draw.polygon(
            [(cx - 40, cy - 180), (cx + 80, cy - 180), (cx + 60, cy - 40), (cx - 20, cy - 40)],
            outline=(140, 190, 160),
        )
        draw.ellipse([cx + 40, cy + 40, cx + 160, cy + 160], outline=(140, 190, 160), width=4)
    elif theme == "ambulance":
        try:
            draw.rounded_rectangle(
                [cx - 360, cy - 80, cx + 360, cy + 160],
                radius=40,
                outline=(200, 80, 80),
                width=5,
            )
        except Exception:
            draw.rectangle([cx - 360, cy - 80, cx + 360, cy + 160], outline=(200, 80, 80), width=5)
        draw.rectangle([cx + 80, cy - 80, cx + 360, cy + 40], outline=(200, 80, 80), width=4)
        draw.ellipse([cx - 220, cy + 120, cx - 100, cy + 240], outline=(200, 120, 120), width=4)
        draw.ellipse([cx + 100, cy + 120, cx + 220, cy + 240], outline=(200, 120, 120), width=4)
    else:
        _draw_ecg_wave(draw, cy, (80, 200, 180))

    label_font = load_jp_font(42, bold=True)
    # 場面ラベル（病棟・経過など）は表示しない
    _ = label
    _ = label_font
    return img.convert("RGBA")


THEME_KEYWORDS: dict[str, list[str]] = {
    "er": ["救急", "ER", "ショック", "心肺停止", "CPA", "挿管", "外傷", "一次評価", "トリアージ"],
    "icu": ["ICU", "集中治療", "モニター", "人工呼吸", "昇圧", "敗血症", "鎮静", "カテコラミン"],
    "surgery": ["手術", "オペ", "開胸", "開腹", "麻酔", "執刀", "縫合", "ドレーン"],
    "lab": ["検査", "採血", "CT", "MRI", "レントゲン", "培養", "病理", "エコー", "画像"],
    "ward": ["病棟", "回診", "退院", "入院", "経過", "ナース", "ベッド"],
    "consult": ["説明", "同意", "家族", "インフォームド", "外来", "決断", "選択"],
    "pharma": ["投与", "点滴", "抗生", "抗菌", "薬", "輸液", "ステロイド", "抗凝固"],
    "ambulance": ["救急隊", "救急車", "現場", "ドクターカー", "搬送", "通報"],
}
THEME_CYCLE = ["er", "icu", "lab", "consult", "pharma", "surgery", "ward", "ambulance"]


def infer_theme_from_text(segment: str, index: int) -> str:
    text = segment or ""
    scores = {k: 0 for k in THEME_KEYWORDS}
    for theme, words in THEME_KEYWORDS.items():
        for word in words:
            if word in text:
                scores[theme] += 1
    best = max(scores, key=lambda k: scores[k])
    if scores[best] == 0:
        return THEME_CYCLE[index % len(THEME_CYCLE)]
    return best


def split_script_for_scenes(script: str, n_scenes: int) -> list[str]:
    text = (script or "").strip()
    if n_scenes <= 1:
        return [text]
    if not text:
        return [""] * n_scenes
    paras = [p for p in re.split(r"\n+", text) if p.strip()]
    if len(paras) >= n_scenes:
        buckets = [""] * n_scenes
        for i, p in enumerate(paras):
            idx = min(i * n_scenes // len(paras), n_scenes - 1)
            buckets[idx] += (("\n" if buckets[idx] else "") + p)
        return buckets
    size = max(1, len(text) // n_scenes)
    parts = []
    for i in range(n_scenes):
        start = i * size
        end = len(text) if i == n_scenes - 1 else (i + 1) * size
        parts.append(text[start:end])
    return parts



def make_fallback_landscape(index: int) -> Image.Image:
    """ダウンロード失敗時は自作の医療イメージ（写真不使用）を使う。"""
    theme = THEME_CYCLE[index % len(THEME_CYCLE)]
    return draw_medical_scene(theme).convert("RGB")


def ensure_landscape_images(needed: int | list[int]) -> list[Path]:
    """
    著作権フリーの医療背景写真を取得してキャッシュする。
    needed が整数なら 0..(needed-1)、リストならその番号の写真を用意。
    失敗時は自作の医療イメージ。
    """
    MEDICAL_BG_DIR.mkdir(parents=True, exist_ok=True)
    urls = MEDICAL_BACKGROUND_URLS
    n_urls = max(1, len(urls))
    if isinstance(needed, int):
        indices = list(range(max(1, int(needed))))
    else:
        indices = [int(x) for x in needed] or [0]

    unique = sorted({i % n_urls for i in indices})
    path_by_idx: dict[int, Path] = {}
    for i in unique:
        out = MEDICAL_BG_DIR / f"med_bg_{i:03d}.jpg"
        if out.exists() and out.stat().st_size > 8000:
            path_by_idx[i] = out
            continue
        url = urls[i % n_urls]
        ok = False
        try:
            r = http_session_direct().get(url, timeout=60, allow_redirects=True)
            if r.status_code == 200 and len(r.content) > 8000:
                ctype = (r.headers.get("Content-Type") or "").lower()
                if "html" not in ctype:
                    with Image.open(io.BytesIO(r.content)) as im:
                        im.load()
                    out.write_bytes(r.content)
                    ok = True
        except Exception:
            ok = False
        if not ok:
            img = make_fallback_landscape(i)
            img.save(out, format="JPEG", quality=90)
        path_by_idx[i] = out

    return [path_by_idx[i % n_urls] for i in indices]


def plan_scene_schedule(script: str, total_duration: float) -> list[dict[str, Any]]:
    """約1分ごとに医療背景を切り替えるスケジュール。"""
    total_duration = max(float(total_duration), 1.0)
    n = max(1, int((total_duration + SCENE_INTERVAL_SEC - 0.01) // SCENE_INTERVAL_SEC))
    segments = split_script_for_scenes(script, n)
    n_urls = max(1, len(MEDICAL_BACKGROUND_URLS))
    schedule = []
    prev_bg: int | None = None
    for i in range(n):
        start = i * SCENE_INTERVAL_SEC
        end = min((i + 1) * SCENE_INTERVAL_SEC, total_duration)
        dur = max(0.5, end - start)
        seg = segments[i] if i < len(segments) else ""
        theme = infer_theme_from_text(seg, i)
        bg = int(THEME_TO_BG_INDEX.get(theme, i % n_urls)) % n_urls
        if prev_bg is not None and bg == prev_bg:
            bg = (bg + 1) % n_urls
        prev_bg = bg
        schedule.append(
            {
                "index": i,
                "theme": theme,
                "landscape_index": bg,
                "duration": dur,
                "segment": seg,
            }
        )
    spent = sum(s["duration"] for s in schedule[:-1]) if len(schedule) > 1 else 0
    if schedule:
        schedule[-1]["duration"] = max(0.5, total_duration - spent)
    return schedule



def create_scene_frame(
    path: Path,
    landscape_path: Path | None = None,
    title: str = "",
    title_img: Image.Image | None = None,
    show_title: bool = False,
    disclaimer: str = DISCLAIMER_TEXT,
    landscape_index: int = 0,
) -> Path:
    """
    医療写真背景＋注意書き。
    本編にはタイトル文字・VOICEVOXクレジット・場面ラベルを出さない
    （タイトルはタイトル画像側で扱う）。
    """
    w, h = VIDEO_SIZE
    if landscape_path is not None and Path(landscape_path).exists():
        try:
            photo = Image.open(landscape_path).convert("RGB")
            base = fit_image_cover(photo, VIDEO_SIZE).convert("RGBA")
        except Exception:
            base = make_fallback_landscape(landscape_index).convert("RGBA")
    else:
        base = make_fallback_landscape(landscape_index).convert("RGBA")

    # 下部を少し暗くして注意書きを読みやすく
    shade = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    sd = ImageDraw.Draw(shade)
    sd.rectangle([0, int(h * 0.82), w, h], fill=(0, 0, 0, 140))
    if show_title:
        sd.rectangle([0, 0, w, int(h * 0.28)], fill=(0, 0, 0, 70))
    base = Image.alpha_composite(base, shade)
    draw = ImageDraw.Draw(base)

    if show_title:
        if title_img is not None:
            paste_centered(
                base,
                title_img,
                max_width_ratio=0.92,
                max_height_ratio=0.28,
                y_ratio=0.06,
            )
            draw = ImageDraw.Draw(base)
        else:
            max_text_w = int(w * 0.88)
            title = (title or "").strip() or "医学ドラマ"
            title_font = load_jp_font(88, bold=True)
            pink = (255, 80, 160)
            y_cursor = int(h * 0.08)
            for line in wrap_text_to_width(title, title_font, max_text_w, draw)[:3]:
                bbox = draw.textbbox((0, 0), line, font=title_font)
                tw = bbox[2] - bbox[0]
                th = bbox[3] - bbox[1]
                draw_outlined_text(
                    draw,
                    ((w - tw) // 2, y_cursor),
                    line,
                    title_font,
                    fill=pink,
                    outline=(0, 0, 0),
                    outline_width=8,
                )
                y_cursor += th + 14

    # 注意書きのみ（一番下）。VOICEVOXクレジットは本編に出さない
    margin = 36
    disc_font = load_jp_font(22, bold=False)
    disc = (disclaimer or DISCLAIMER_TEXT).strip()
    disc_lines = wrap_text_to_width(disc, disc_font, int(w * 0.94), draw)
    y_disc = h - margin
    for line in reversed(disc_lines[-3:]):
        bbox = draw.textbbox((0, 0), line, font=disc_font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        y_disc -= th + 4
        x = (w - tw) // 2
        draw.text((x + 1, y_disc + 1), line, font=disc_font, fill=(0, 0, 0))
        draw.text((x, y_disc), line, font=disc_font, fill=(210, 215, 220))

    path.parent.mkdir(parents=True, exist_ok=True)
    base.convert("RGB").save(path, format="PNG")
    return path


def create_background_image(
    path: Path,
    title: str = "",
    subtitle: str = "",
    background_img: Image.Image | None = None,
    title_img: Image.Image | None = None,
    footnote_text: str = "",
    voicevox_credit: str = CREDIT_TEXT,
) -> Path:
    _ = background_img
    _ = footnote_text
    _ = subtitle
    _ = voicevox_credit
    return create_scene_frame(
        path,
        landscape_path=None,
        title=title,
        title_img=title_img,
        show_title=False,
        landscape_index=0,
    )



# ---------------------------------------------------------------------------
# BGM（現在は未使用：動画にはナレーションのみ。必要になったら再利用可）
# ---------------------------------------------------------------------------
def make_fallback_bgm_wav(path: Path, seconds: float = 60.0, rate: int = 44100) -> Path:
    """ダウンロード失敗時: 低いドローン音のWAVを自作（追加ライブラリ不要）。"""
    import math

    path = path.with_suffix(".wav")
    n = int(seconds * rate)
    frames = bytearray()
    for i in range(n):
        t = i / rate
        # 低いサイン波を重ねてシリアスな雰囲気に
        v = 0.15 * math.sin(2 * math.pi * 55 * t)
        v += 0.08 * math.sin(2 * math.pi * 82.5 * t)
        v += 0.05 * math.sin(2 * math.pi * 110 * t)
        # フェードイン/アウト
        fade = min(t / 2.0, 1.0, max(0.0, (seconds - t) / 2.0))
        sample = int(max(-1.0, min(1.0, v * fade)) * 32767)
        frames += struct.pack("<h", sample)

    with wave.open(str(path), "wb") as w:
        w.setnchannels(1)
        w.setsampwidth(2)
        w.setframerate(rate)
        w.writeframes(frames)
    return path


def ensure_bgm(work_dir: Path = WORK_DIR) -> Path:
    """bgm.mp3 が無ければ取得。失敗時はWAVを生成。"""
    mp3_path = work_dir / BGM_FILENAME
    if mp3_path.exists() and mp3_path.stat().st_size > 1000:
        return mp3_path

    for url in BGM_CANDIDATE_URLS:
        try:
            r = http_session_direct().get(url, timeout=60, allow_redirects=True)
            if r.status_code == 200 and len(r.content) > 1000:
                ctype = r.headers.get("Content-Type", "")
                if "html" in ctype.lower():
                    continue
                mp3_path.write_bytes(r.content)
                return mp3_path
        except requests.RequestException:
            continue

    return make_fallback_bgm_wav(work_dir / "bgm_fallback")


# ---------------------------------------------------------------------------
# 動画合成（moviepy）
# ---------------------------------------------------------------------------
def build_mp4(
    narration_wav: Path,
    scene_clips: list[tuple[Path, float]],
    output_mp4: Path,
    ending_png: Path | None = None,
    ending_duration: float = ENDING_DURATION_SEC,
    subtitle_cues: list[dict[str, Any]] | None = None,
    subtitle_dir: Path | None = None,
    progress_callback: Callable[[float], None] | None = None,
) -> Path:
    """
    医療静止画シーン + 長尺音声 + 同期字幕 + エンディング著作権表示。
    BGM・効果音は入れない（ナレーション音声のみ）。
    scene_clips: [(画像パス, 秒数), ...]
    subtitle_cues: [{start, end, text}, ...]  textはルビなし
    """
    from moviepy.editor import (
        AudioFileClip,
        CompositeVideoClip,
        ImageClip,
        concatenate_videoclips,
    )

    voice = AudioFileClip(str(narration_wav))
    duration = float(voice.duration)
    if duration <= 0:
        voice.close()
        raise RuntimeError("音声の長さが 0 です。")
    if not scene_clips:
        voice.close()
        raise RuntimeError("シーン画像がありません。")

    # moviepy の音声長と、wave 計測の差を字幕全体に同じ倍率で吸収する
    # （最後の字幕終了時刻ではなく、WAV全体の長さを基準にする＝同期を崩さない）
    if subtitle_cues:
        with wave.open(str(narration_wav), "rb") as wf:
            wave_dur = wf.getnframes() / float(wf.getframerate())
        if wave_dur > 0.5 and abs(wave_dur - duration) > 0.01:
            scale = duration / wave_dur
            for c in subtitle_cues:
                c["start"] = float(c.get("start", 0)) * scale
                c["end"] = float(c.get("end", 0)) * scale

    # 字幕は音声の実時間どおりに出す（フレーム丸めはしない＝ずれ防止）
    still_fps = SUBTITLE_VIDEO_FPS
    clips = []
    for img_path, dur in scene_clips:
        clips.append(
            ImageClip(str(img_path)).set_duration(float(dur)).set_fps(still_fps)
        )
    main_video = concatenate_videoclips(clips, method="compose")
    # シーン合計と音声長さの差を吸収（字幕時刻は音声基準）
    if abs(float(main_video.duration) - duration) > 0.05:
        main_video = main_video.set_duration(duration)

    sub_clips = []
    if subtitle_cues:
        sub_dir = subtitle_dir or (output_mp4.parent / "_subs")
        sub_dir.mkdir(parents=True, exist_ok=True)
        for i, cue in enumerate(subtitle_cues):
            start = float(cue.get("start", 0))
            end = float(cue.get("end", 0))
            text = strip_voicevox_ruby(str(cue.get("text") or "")).strip()
            if not text or end <= start:
                continue
            if start >= duration:
                continue
            end = min(end, duration)
            # 最低でも1フレーム分は出す（ただし開始は音声どおり）
            min_dur = 1.0 / float(still_fps)
            if end - start < min_dur:
                end = min(duration, start + min_dur)
            png = sub_dir / f"sub_{i:05d}.png"
            create_subtitle_png(text, png)
            sub_clips.append(
                ImageClip(str(png), ismask=False)
                .set_start(start)
                .set_duration(max(min_dur, end - start))
                .set_fps(still_fps)
                .set_position((0, 0))
            )

    if sub_clips:
        main_video = CompositeVideoClip(
            [main_video] + sub_clips, size=VIDEO_SIZE
        ).set_duration(duration)

    # ナレーションのみ（BGM・効果音なし）
    main_video = main_video.set_audio(voice)

    parts = [main_video]
    ending_clip = None
    transition_clip = None
    ending_hold_clip = None
    if ending_png is not None and ending_duration > 0:
        fade_sec = float(ENDING_FADE_SEC)
        end_dur = float(ending_duration)
        ending_base = (
            ImageClip(str(ending_png))
            .set_fps(still_fps)
        )
        # 本編音声終了後: 最終フレーム → エンディングへ 5秒でフェード移行
        last_t = max(0.0, duration - (1.0 / float(still_fps)))
        last_frame = (
            main_video.to_ImageClip(t=last_t)
            .set_duration(fade_sec)
            .set_fps(still_fps)
            .fadeout(fade_sec)
        )
        ending_fade_in = (
            ending_base.set_duration(fade_sec)
            .fadein(fade_sec)
        )
        transition_clip = CompositeVideoClip(
            [last_frame, ending_fade_in],
            size=VIDEO_SIZE,
        ).set_duration(fade_sec).set_fps(still_fps)
        # フェード完了後、エンディングをそのまま表示
        ending_hold_clip = ending_base.set_duration(end_dur).set_fps(still_fps)
        ending_clip = transition_clip  # close用の代表
        parts.extend([transition_clip, ending_hold_clip])

    video = concatenate_videoclips(parts, method="compose")

    output_mp4.parent.mkdir(parents=True, exist_ok=True)
    tried_errors: list[str] = []
    encode_attempts = [
        {
            "codec": "h264_videotoolbox",
            "audio_codec": "aac",
            "ffmpeg_params": ["-b:v", "3M", "-pix_fmt", "yuv420p"],
        },
        {
            "codec": "libx264",
            "audio_codec": "aac",
            "preset": "ultrafast",
            "ffmpeg_params": ["-pix_fmt", "yuv420p", "-tune", "stillimage"],
        },
    ]

    last_err: Exception | None = None
    encode_logger = None
    if progress_callback is not None:
        try:
            from proglog import ProgressBarLogger

            class _EncodeProgressLogger(ProgressBarLogger):
                def bars_callback(self, bar, attr, value, old_value=None):
                    if bar != "t" or attr != "index":
                        return
                    total = (self.bars.get("t") or {}).get("total") or 0
                    if total:
                        progress_callback(float(value) / float(total))

            encode_logger = _EncodeProgressLogger()
        except Exception:
            encode_logger = None

    for opts in encode_attempts:
        try:
            kwargs = {
                "fps": still_fps,
                "codec": opts["codec"],
                "audio_codec": opts["audio_codec"],
                "threads": 0,
                "logger": encode_logger,
            }
            if "preset" in opts:
                kwargs["preset"] = opts["preset"]
            if "ffmpeg_params" in opts:
                kwargs["ffmpeg_params"] = opts["ffmpeg_params"]
            video.write_videofile(str(output_mp4), **kwargs)
            last_err = None
            break
        except Exception as e:  # noqa: BLE001
            tried_errors.append(f"{opts['codec']}: {e}")
            last_err = e
            if output_mp4.exists():
                try:
                    output_mp4.unlink()
                except OSError:
                    pass

    voice.close()
    video.close()
    try:
        main_video.close()
    except Exception:
        pass
    if ending_clip is not None:
        try:
            ending_clip.close()
        except Exception:
            pass
    if transition_clip is not None and transition_clip is not ending_clip:
        try:
            transition_clip.close()
        except Exception:
            pass
    if ending_hold_clip is not None:
        try:
            ending_hold_clip.close()
        except Exception:
            pass
    for c in clips:
        try:
            c.close()
        except Exception:
            pass
    for c in sub_clips:
        try:
            c.close()
        except Exception:
            pass

    if last_err is not None:
        raise RuntimeError(
            "動画エンコードに失敗しました。\n" + "\n".join(tried_errors)
        ) from last_err

    return output_mp4


# ---------------------------------------------------------------------------
# UI ヘルパー（レビュー採否）
# ---------------------------------------------------------------------------
REVIEW_SECTION_DEFS = [
    ("medical_contradictions", "(1) 医学的に矛盾している箇所"),
    ("awkward_for_doctors", "(2) 現役の医師が聞くと違和感がある表現"),
    ("immersion_improvements", "(3) 修正すると臨場感が増す箇所"),
]

CHOICE_ACCEPT = "accept"
CHOICE_REJECT = "reject"
CHOICE_REVISE = "revise"
CHOICE_LABELS = {
    CHOICE_ACCEPT: "①承諾（そのまま反映）",
    CHOICE_REJECT: "②却下",
    CHOICE_REVISE: "③別案にて修正",
}


def decision_widget_key(section_key: str, index: int) -> str:
    return f"review_choice__{section_key}__{index}"


def alt_widget_key(section_key: str, index: int) -> str:
    return f"review_alt__{section_key}__{index}"


def clear_review_decision_widgets(review: dict[str, Any] | None) -> None:
    """新しいレビュー結果に合わせて、古い採否ウィジェット状態を消す。"""
    if not review:
        return
    for section_key, _ in REVIEW_SECTION_DEFS:
        items = review.get(section_key, []) or []
        for i in range(len(items)):
            for k in (decision_widget_key(section_key, i), alt_widget_key(section_key, i)):
                if k in st.session_state:
                    del st.session_state[k]


def render_review_section_interactive(
    section_key: str, title: str, items: list[dict[str, str]]
) -> None:
    """各指摘に 承諾／却下／別案 を選べるUI。"""
    st.subheader(title)
    if not items:
        st.caption("該当なし")
        return

    for i, item in enumerate(items):
        label = item.get("original") or "（箇所）"
        with st.expander(f"{i + 1}. {label}", expanded=(i == 0)):
            st.write(item.get("issue") or "（なし）")
            st.write(item.get("suggestion") or "（なし）")

            choice_key = decision_widget_key(section_key, i)
            if choice_key not in st.session_state:
                st.session_state[choice_key] = CHOICE_REJECT

            st.radio(
                "対応",
                options=[CHOICE_ACCEPT, CHOICE_REJECT, CHOICE_REVISE],
                format_func=lambda x: CHOICE_LABELS.get(x, x),
                key=choice_key,
                horizontal=True,
            )

            if st.session_state.get(choice_key) == CHOICE_ACCEPT:
                preview = clean_script_replacement_text(
                    item.get("suggestion") or "", item.get("original") or ""
                )
                if preview:
                    st.caption(f"反映文: {preview}")
                else:
                    st.warning("自動反映不可。「別案」で本文を書いてください。")

            if st.session_state.get(choice_key) == CHOICE_REVISE:
                alt_key = alt_widget_key(section_key, i)
                if alt_key not in st.session_state:
                    st.session_state[alt_key] = item.get("suggestion") or ""
                st.text_area(
                    "別案",
                    key=alt_key,
                    height=100,
                )


def clean_script_replacement_text(suggestion: str, original: str = "") -> str:
    """
    レビュー修正案から、台本へ入れる本文だけを取り出す。
    「編集メモを削除する」「確定文にして」などの解説・手順は捨てる。
    """
    text = (suggestion or "").strip()
    if not text:
        return ""

    # 「」『』内の本文を優先（解説付き提案でよく使われる）
    quoted = re.findall(r"[「『]([^」』]+)[」』]", text)
    quoted = [q.strip() for q in quoted if q.strip()]
    meta_hint = re.compile(
        r"(編集メモ|編集注|編集コメント|確定文|地の文|削除する|してください|採用し)"
    )
    usable_quotes = [q for q in quoted if not meta_hint.search(q)]
    if usable_quotes:
        text = max(usable_quotes, key=len)

    # 末尾〜文中の作業指示を除去
    strip_patterns = [
        r"[、,]?\s*と確定文にして編集メモを削除する。?",
        r"[、,]?\s*と確定文に書き直し[、,]?編集メモを削除する。?",
        r"[、,]?\s*と確定文に書き直す。?",
        r"[、,]?\s*と確定文にして。?",
        r"[、,]?\s*のみを地の文として採用し[、,]?編集注を削除する。?",
        r"[、,]?\s*を地の文として採用し[、,]?編集注を削除する。?",
        r"[、,]?\s*地の文として採用し[、,]?編集注を削除する。?",
        r"編集メモを削除する。?",
        r"編集注を削除する。?",
        r"編集コメントを削除する。?",
        r"[。．]?[、,]?\s*編集(?:メモ|注|コメント).*$",
        r"[。．]?[、,]?\s*確定文に.*$",
        r"[。．]?[、,]?\s*地の文として.*$",
        r"に修正してください。?",
        r"に言い換えてください。?",
        r"に直してください。?",
        r"を推奨します。?",
        r"が自然です。?",
        r"がよいです。?",
        r"など具体的に。?",
        r"など具体値を1か所入れてください。?",
        r"してください。?",
    ]
    for pat in strip_patterns:
        text = re.sub(pat, "", text)

    text = text.strip(" 　\n\r\t「」『』\"'、,")

    # まだ作業指示だけの文章なら空にする（自動反映しない）
    if re.search(
        r"(編集メモ|編集注|確定文|地の文として|削除する|してください|手修正)",
        text,
    ):
        # 引用抽出に失敗し、指示文が残っている
        if original and original in text and len(text) > len(original) + 10:
            # 原文の後に指示が続く場合は原文側だけ残さない（危険なので空）
            return ""
        if not usable_quotes:
            return ""

    return text.strip()


def apply_review_decisions_to_script(
    script: str, review: dict[str, Any]
) -> tuple[str, list[str], list[str]]:
    """
    採択／別案を台本に反映する。
    戻り値: (新しい台本, 反映できた一覧, 手動編集が必要な一覧)
    """
    text = script
    applied: list[str] = []
    manual: list[str] = []

    jobs: list[tuple[int, str, str, str]] = []
    for section_key, section_title in REVIEW_SECTION_DEFS:
        items = review.get(section_key, []) or []
        for i, item in enumerate(items):
            choice = st.session_state.get(
                decision_widget_key(section_key, i), CHOICE_REJECT
            )
            if choice == CHOICE_REJECT:
                continue

            original = (item.get("original") or "").strip()
            suggestion = (item.get("suggestion") or "").strip()
            if choice == CHOICE_ACCEPT:
                replacement = clean_script_replacement_text(suggestion, original)
            else:
                raw_alt = (
                    st.session_state.get(alt_widget_key(section_key, i), suggestion)
                    or ""
                ).strip()
                # 別案も解説文が混ざっていたら除去（ユーザーが書いた文はできるだけ残す）
                replacement = clean_script_replacement_text(raw_alt, original)
                if not replacement and raw_alt and not re.search(
                    r"(編集メモ|編集注|確定文|地の文として採用)", raw_alt
                ):
                    replacement = raw_alt

            label = f"{section_title} #{i + 1}"
            if not original or not replacement:
                manual.append(
                    f"{label}: 差し替え本文を取り出せませんでした"
                    f"（修正案に解説だけがある可能性があります。手修正してください）\n"
                    f"→ 元の修正案: {suggestion}"
                )
                continue
            if original in ("（全体）", "（バイタル表記なし）", "（箇所）"):
                manual.append(
                    f"{label}: 全体向けの指摘のため自動反映できません"
                    f"（別案/修正案を手で入れてください）\n→ {replacement}"
                )
                continue
            pos = text.find(original)
            if pos < 0:
                manual.append(
                    f"{label}: 台本内に『{original}』が見つかりません"
                    f"（手修正してください）\n→ {replacement}"
                )
                continue
            jobs.append((pos, original, replacement, label))

    jobs.sort(key=lambda x: x[0], reverse=True)
    for pos, original, replacement, label in jobs:
        if text[pos : pos + len(original)] != original:
            pos2 = text.find(original)
            if pos2 < 0:
                manual.append(f"{label}: 反映中に原文が見つからなくなりました")
                continue
            pos = pos2
        text = text[:pos] + replacement + text[pos + len(original) :]
        applied.append(f"{label}: 『{original}』→『{replacement}』")

    return text, applied, manual


def init_state() -> None:
    defaults = {
        "raw_script": "",
        "review": None,
        "final_script": "",
        "review_done": False,
        "script_confirmed": False,
        "ruby_ready": False,
        "ruby_script": "",
        "ruby_skipped": False,
        "skip_review": False,
        "mp4_bytes": None,
        "mp4_path": "",
        "mp4_name": "medical_drama.mp4",
        "last_error": "",
        "video_title": "",
        "title_suggestion": "",
        "title_decision": "",
        "ruby_dict_custom": None,
        "ruby_dict_source_name": "",
        "ruby_dict_learned": [],
        "ruby_dict_export_ready": False,
        "ruby_dict_last_updates": [],
        "ruby_dict_post_video_updates": [],
        "export_progress_pct": 0,
        "export_progress_msg": "",
        "ruby_script_baseline": "",
        "ending_credits_text": "",
        "reference_text": "",
        "last_script_path": "",
        "last_script_name": "medical_drama.docx",
        "video_encoding": False,
        "_export_job": None,
        "review_apply_log": [],
        "review_manual_log": [],
        "vvox_speaker_name": DEFAULT_SPEAKER_NAME,
        "vvox_style_name": DEFAULT_STYLE_NAME,
        "vvox_style_id": DEFAULT_SPEAKER_ID,
        "vvox_speed_scale": VOICEVOX_SPEED_SCALE,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # 参考文献: 空なら前回保存分を復元（上書きも可）
    if not (st.session_state.get("reference_text") or "").strip():
        saved_ref = load_saved_reference_text()
        if saved_ref:
            st.session_state.reference_text = saved_ref

    # 読み上げ速度を 0.8〜1.5（初期 1.0）にそろえる
    st.session_state.vvox_speed_scale = clamp_voicevox_speed(
        st.session_state.get("vvox_speed_scale", VOICEVOX_SPEED_SCALE)
    )
    # 旧初期値 1.1 / 1.2 のまま残っている場合は、新しい初期値 1.0 に一度だけ更新
    if st.session_state.get("_vvox_speed_default_v10") is not True:
        cur = float(st.session_state.vvox_speed_scale)
        if abs(cur - 1.1) < 1e-9 or abs(cur - 1.2) < 1e-9:
            st.session_state.vvox_speed_scale = float(VOICEVOX_SPEED_SCALE)
        st.session_state["_vvox_speed_default_v10"] = True

    # 声優初期値の移行（旧・No.7 → 青山龍星 ノーマル）を一度だけ
    if st.session_state.get("_vvox_default_v3") is not True:
        old_name = st.session_state.get("vvox_speaker_name")
        if old_name in (None, "", "No.7"):
            st.session_state.vvox_speaker_name = DEFAULT_SPEAKER_NAME
            st.session_state.vvox_style_name = DEFAULT_STYLE_NAME
            st.session_state.vvox_style_id = DEFAULT_SPEAKER_ID
        st.session_state["_vvox_default_v3"] = True


# ---------------------------------------------------------------------------
# Streamlit メイン
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# 動画書き出し
# ---------------------------------------------------------------------------
def run_video_export(progress, pct_box, status) -> None:
    """
    音声・背景・字幕・エンディングをまとめて MP4 にする。
    progress / pct_box / status は Streamlit の表示用オブジェクト。
    """
    def _pct(n: int, msg: str = "") -> None:
        update_export_progress(progress, pct_box, status, n, msg)

    ok, ver = check_voicevox()
    if not ok:
        raise RuntimeError(
            "VOICEVOX に接続できません。アプリを起動してから再実行してください。"
            f"（詳細: {ver}）"
        )

    style_id = int(st.session_state.get("vvox_style_id", DEFAULT_SPEAKER_ID))
    speaker_name = str(
        st.session_state.get("vvox_speaker_name", DEFAULT_SPEAKER_NAME)
    )
    style_name = str(st.session_state.get("vvox_style_name", DEFAULT_STYLE_NAME))
    ending_body = (st.session_state.get("ending_credits_text") or "").strip()
    if not ending_body:
        ending_body = build_ending_credits_text(
            st.session_state.get("reference_text", ""),
            speaker_name,
        )
        st.session_state.ending_credits_text = ending_body
        st.session_state._ending_auto_text = ending_body

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    speed_scale = clamp_voicevox_speed(
        st.session_state.get("vvox_speed_scale", VOICEVOX_SPEED_SCALE)
    )
    _pct(2, "台本を準備中…")
    # 直前工程で確定したルビ入り原稿を使う（ここで辞書ルビを付け直さない）
    voice_script = canonicalize_voicevox_ruby_delimiters(
        str(st.session_state.get("ruby_script") or st.session_state.get("final_script") or "")
    ).strip()
    if not voice_script:
        raise RuntimeError("ルビ入り原稿が空です。先にルビ準備を完了してください。")
    ruby_count = count_voicevox_ruby(voice_script)

    save_reference_text(st.session_state.get("reference_text", ""))

    video_title = str(st.session_state.get("video_title") or "").strip()
    script_docx_name = make_script_docx_filename(video_title)
    script_path = OUTPUT_DIR / "last_script.txt"
    script_path.write_text(voice_script, encoding="utf-8")
    script_docx_path = OUTPUT_DIR / script_docx_name
    script_docx_path.write_bytes(text_to_docx_bytes(voice_script))
    # 旧固定名も残す（互換）
    (OUTPUT_DIR / "last_script.docx").write_bytes(script_docx_path.read_bytes())
    tts_script = expand_voicevox_ruby_to_reading(voice_script)
    sub_script = strip_voicevox_ruby(voice_script)
    (OUTPUT_DIR / "last_script_tts.txt").write_text(tts_script, encoding="utf-8")
    (OUTPUT_DIR / "last_script_subtitle.txt").write_text(
        sub_script, encoding="utf-8"
    )
    st.session_state.last_script_path = str(script_docx_path)
    st.session_state.last_script_name = script_docx_name

    with tempfile.TemporaryDirectory(prefix="meddrama_") as tmp:
        tmp_path = Path(tmp)
        ruby_msg = f"ルビ{ruby_count}件"
        _pct(
            5,
            f"音声生成中（{speaker_name} / {style_name}・{ruby_msg}・"
            f"{speed_scale:.1f}倍）…",
        )
        wav_path = tmp_path / "narration.wav"
        n_chunks = len(split_text_for_voicevox(voice_script))

        def _voice_prog(done: int, total: int) -> None:
            pct = 5 + int(45 * (done / max(total, 1)))
            _pct(
                min(pct, 50),
                f"読み上げ中 {done}/{total}"
                f"（予定 {n_chunks}・{speaker_name} / {style_name}）",
            )

        wav_path, subtitle_cues = generate_narration_wav_to_file(
            voice_script,
            wav_path,
            progress_callback=_voice_prog,
            speaker=style_id,
            speed_scale=speed_scale,
        )
        # 同期確認は無音を足す前（朗読と字幕の対応を検査）
        sync_issues = validate_audio_subtitle_sync(wav_path, subtitle_cues)
        if sync_issues:
            raise RuntimeError(
                "音声と字幕の同期チェックに失敗しました:\n"
                + "\n".join(f"- {m}" for m in sync_issues)
            )

        # 冒頭に無音を入れ、字幕も同じ秒数だけ後ろへずらす
        intro_silence = float(INTRO_SILENCE_SEC)
        if intro_silence > 0:
            _pct(51, f"冒頭に {intro_silence:.0f} 秒の無音を追加…")
            prepend_silence_to_wav(wav_path, intro_silence)
            shift_subtitle_cues(subtitle_cues, intro_silence)

        import wave as _wave

        with _wave.open(str(wav_path), "rb") as wf:
            audio_sec = wf.getnframes() / float(wf.getframerate())

        schedule = plan_scene_schedule(st.session_state.final_script, audio_sec)
        _pct(52, f"背景準備（{len(schedule)} 枚）…")
        bg_indices = [
            int(item.get("landscape_index", item["index"])) for item in schedule
        ]
        landscapes = ensure_landscape_images(bg_indices)
        _pct(55, "シーン合成…")
        scene_dir = tmp_path / "scenes"
        scene_dir.mkdir(parents=True, exist_ok=True)
        scene_clips: list[tuple[Path, float]] = []
        for item in schedule:
            i = int(item["index"])
            li = int(item.get("landscape_index", i))
            dur = float(item["duration"])
            frame_path = scene_dir / f"scene_{i:03d}.png"
            land = landscapes[i % len(landscapes)]
            create_scene_frame(
                frame_path,
                landscape_path=land,
                title="",
                title_img=None,
                show_title=False,
                disclaimer=DISCLAIMER_TEXT,
                landscape_index=li,
            )
            scene_clips.append((frame_path, dur))
            pct = 55 + int(12 * ((i + 1) / max(len(schedule), 1)))
            _pct(min(pct, 67), f"シーン {i+1}/{len(schedule)}")

        _pct(70, "エンディング…")
        ending_path = tmp_path / "ending_credits.png"
        create_ending_credits_frame(ending_path, ending_text=ending_body)
        (OUTPUT_DIR / "last_ending.png").write_bytes(ending_path.read_bytes())
        (OUTPUT_DIR / "last_frame.png").write_bytes(
            scene_clips[0][0].read_bytes()
        )

        _pct(75, "MP4 エンコード中…")

        def _encode_prog(frac: float) -> None:
            # 75%〜99% をエンコード進捗に割り当て
            pct = 75 + int(24 * max(0.0, min(1.0, float(frac))))
            _pct(min(pct, 99), f"MP4 エンコード中… {int(frac * 100)}%")

        out_path = OUTPUT_DIR / "medical_drama.mp4"
        build_mp4(
            wav_path,
            scene_clips,
            out_path,
            ending_png=ending_path,
            ending_duration=ENDING_DURATION_SEC,
            subtitle_cues=subtitle_cues,
            subtitle_dir=tmp_path / "_subs",
            progress_callback=_encode_prog,
        )

        desktop_dir = get_desktop_dir()
        desktop_name = make_desktop_mp4_filename(video_title)
        desktop_path = desktop_dir / desktop_name
        shutil.copy2(out_path, desktop_path)

        st.session_state.mp4_path = str(desktop_path)
        st.session_state.mp4_name = desktop_name
        st.session_state.mp4_bytes = None
        # ルビ入り最終原稿 ↔ 辞書を比較し、足りないルビを追加
        _pct(99, "ルビ辞書を更新中…")
        applied = sync_script_rubies_into_dictionary(voice_script)
        if applied:
            _pct(
                100,
                f"完了（辞書にルビを {len(applied)} 件追加）",
            )
        else:
            _pct(100, "完了")
        status.success(f"完了: {desktop_path}")


def inject_app_theme() -> None:
    """余計な装飾を抑え、読みやすい白背景にする。"""
    st.markdown(
        """
<style>
  [data-testid="stAppViewContainer"] { background: #fff; color: #111; }
  [data-testid="stHeader"] { background: #fff; }
  [data-testid="stSidebar"] { background: #fafafa; }
  h1, h2, h3, h4 { font-size: 1rem !important; font-weight: 600 !important; }
  div[data-testid="stAlert"] { border: 1px solid #ccc !important; }
</style>
        """,
        unsafe_allow_html=True,
    )


def main() -> None:
    st.set_page_config(
        page_title="医学ドラマ動画メーカー",
        page_icon=None,
        layout="wide",
    )
    inject_app_theme()
    init_state()

    # MP4作成中は他UIを出さず、誤操作を防ぐ
    # Stop／再読み込みで中断されたあとも通常画面に戻れるようにする
    if st.session_state.get("video_encoding"):
        st.write("医学ドラマ動画メーカー")
        st.warning("動画作成中です。完了するまでこのページを閉じないでください。")
        st.markdown(
            '<div style="font-size:1.2rem;margin-bottom:0.5rem;">'
            "作業の進捗（パーセント）</div>",
            unsafe_allow_html=True,
        )
        if st.button("中止して通常画面に戻る", key="btn_cancel_video_encoding"):
            st.session_state.video_encoding = False
            st.session_state._export_job = None
            st.rerun()

        job = st.session_state.get("_export_job")
        # pending / running 以外（古い状態・不明）は自動再開せず中断扱いにする
        if job not in ("pending", "running"):
            job = "running"
        # Stop などで中断されたあと：自動再開せず、再開／中止を選ばせる
        if job == "running":
            st.error("前回の作成が中断されたか、作成モードのまま残っています。")
            if st.button("最初から再開する", type="primary", key="btn_restart_export"):
                st.session_state._export_job = "pending"
                st.rerun()
            st.stop()

        # pending → 書き出し開始
        st.session_state._export_job = "running"
        st.session_state.export_progress_pct = 0
        st.session_state.export_progress_msg = "準備中…"
        progress = st.progress(0, text="0%  準備中…")
        pct_box = st.empty()
        status = st.empty()
        update_export_progress(progress, pct_box, status, 0, "準備中…")
        try:
            run_video_export(progress, pct_box, status)
        except Exception as e:  # noqa: BLE001
            st.session_state.mp4_path = ""
            st.session_state.mp4_bytes = None
            st.session_state.video_encoding = False
            st.session_state._export_job = None
            st.error(f"動画生成に失敗しました: {e}")
            st.exception(e)
            st.stop()
        st.session_state.video_encoding = False
        st.session_state._export_job = None
        st.rerun()

    st.write("医学ドラマ動画メーカー")

    with st.sidebar:
        st.write("設定")
        ok, ver = check_voicevox()
        if ok:
            st.caption(f"VOICEVOX OK（{ver}）")
        else:
            st.error(f"VOICEVOX 未接続: {ver}")

        load_dotenv_file()
        has_saved = bool(get_api_key())
        typed = st.text_input(
            "APIキー",
            type="password",
            placeholder="ANTHROPIC_API_KEY",
            help="保存済み" if has_saved else "未設定",
        )
        if typed.strip():
            os.environ["ANTHROPIC_API_KEY"] = typed.strip()

        if st.button("キーを保存"):
            to_save = typed.strip() or get_api_key()
            if not to_save:
                st.error("先にキーを入力してください。")
            else:
                try:
                    saved_path = save_api_key_to_env_file(to_save)
                    st.success(f"保存: `{saved_path.name}`")
                except Exception as e:  # noqa: BLE001
                    st.error(f"保存失敗: {e}")

    # ----- Step 1 -----
    st.write("1. 台本")

    input_mode = st.radio(
        "入力方法",
        options=["paper", "script"],
        format_func=lambda x: (
            "論文PDFから作る"
            if x == "paper"
            else "台本ファイルを取り込む"
        ),
        key="step1_input_mode",
        horizontal=False,
        label_visibility="collapsed",
    )

    if input_mode == "paper":
        paper_pdf = st.file_uploader(
            "医学論文PDF",
            type=["pdf"],
            key="paper_pdf_upload",
        )
        if st.button(
            "PDFから台本を作成",
            type="primary",
            key="btn_pdf_to_script",
            use_container_width=True,
        ):
            api_key = get_api_key()
            if not api_key:
                st.error("先にAPIキーを入力・保存してください。")
            elif paper_pdf is None:
                st.error("PDFを選んでください。")
            else:
                try:
                    with st.spinner("PDFの文字を読み取っています…"):
                        raw = paper_pdf.getvalue()
                        paper_text = extract_text_from_pdf_bytes(raw)
                    with st.spinner(
                        "台本を作成中です（数分かかることがあります）…"
                    ):
                        script = generate_drama_script_from_paper(paper_text, api_key)
                    if not script.strip():
                        st.error("台本が空でした。別のPDFで試してください。")
                    else:
                        # ルビなし原稿として取り込む（ルビは動画直前に付与）
                        plain = strip_voicevox_ruby(script).strip()
                        commit_loaded_script(
                            plain,
                            f"pdf-{paper_pdf.name}-{len(plain)}",
                        )
                        script = plain
                        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
                        (OUTPUT_DIR / "last_script.txt").write_text(
                            script, encoding="utf-8"
                        )
                        # 論文から Vancouver 方式の参考文献を作り、エンディングへ反映
                        with st.spinner("参考文献を作成中…"):
                            citation = extract_vancouver_citation_from_paper(
                                paper_text, api_key
                            )
                        if citation:
                            apply_paper_reference_to_session(citation)
                        else:
                            apply_paper_reference_to_session(
                                f"（PDFより作成・書誌情報を確認してください）{paper_pdf.name}"
                            )
                        st.success(
                            f"取り込み完了: {paper_pdf.name}"
                            f"（約 {len(script):,} 字・ルビなし）"
                        )
                except Exception as e:  # noqa: BLE001
                    st.error(f"台本作成に失敗しました: {e}")

    else:
        script_upload = st.file_uploader(
            "台本ファイル（.docx / .pdf）",
            type=["docx", "pdf"],
            key="ready_script_upload",
        )
        if st.button(
            "台本を取り込む",
            type="primary",
            key="btn_import_ready_script",
            use_container_width=True,
        ):
            if script_upload is None:
                st.error("台本ファイルを選んでください。")
            else:
                try:
                    with st.spinner("台本を読み取っています…"):
                        raw = script_upload.getvalue()
                        script = extract_text_from_bytes(
                            script_upload.name, raw
                        ).strip()
                    if not script:
                        st.error("台本が空でした。")
                    else:
                        plain = strip_voicevox_ruby(script).strip()
                        commit_loaded_script(
                            plain,
                            f"script-{script_upload.name}-{len(plain)}",
                        )
                        script = plain
                        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
                        (OUTPUT_DIR / "last_script.txt").write_text(
                            script, encoding="utf-8"
                        )
                        st.success(
                            f"取り込み完了: {script_upload.name}"
                            f"（約 {len(script):,} 字・ルビなし）"
                        )
                except Exception as e:  # noqa: BLE001
                    st.error(f"台本の取り込みに失敗しました: {e}")

    # タイトルは手入力のみ（提案しない）。値は上書きするまで保持。
    if st.session_state.get("raw_script") and not st.session_state.get(
        "script_confirmed"
    ):
        render_video_title_input()

    if st.session_state.raw_script:
        n_chars = len(st.session_state.raw_script)
        est_min = max(1, round(n_chars / 320))
        st.caption(f"{n_chars:,} 字 ／ 目安 {est_min} 分（ルビなし原稿）")
        with st.expander("台本（原文・ルビなし）", expanded=False):
            st.text(st.session_state.raw_script)

        col_rev, col_skip = st.columns(2)
        with col_rev:
            review_clicked = st.button(
                "1. レビューする",
                type="primary",
                disabled=not bool(st.session_state.raw_script.strip()),
                use_container_width=True,
                key="btn_do_review",
            )
        with col_skip:
            skip_clicked = st.button(
                "1′. レビューせず進む",
                type="secondary",
                disabled=not bool(st.session_state.raw_script.strip()),
                use_container_width=True,
                key="btn_skip_review",
            )
    else:
        review_clicked = False
        skip_clicked = False

    if review_clicked:
        with st.spinner("台本をレビューしています…"):
            try:
                # 古い採否の選択を消してから新しいレビューを入れる
                clear_review_decision_widgets(st.session_state.get("review"))
                plain = strip_voicevox_ruby(st.session_state.raw_script)
                st.session_state.review = run_script_review(plain)
                st.session_state.review_done = True
                st.session_state.skip_review = False
                st.session_state.script_confirmed = False
                st.session_state.ruby_ready = False
                st.session_state.mp4_bytes = None
                st.session_state.mp4_path = ""
                st.session_state.review_apply_log = []
                st.session_state.review_manual_log = []
                st.session_state.last_error = ""
                # 最終台本はルビなしで確定する
                st.session_state.final_script = plain
                st.session_state.final_script_editor = plain
                st.session_state.final_script_editor_widget = plain
            except Exception as e:  # noqa: BLE001
                st.session_state.last_error = str(e)
                st.error(f"レビューに失敗しました: {e}")

    if skip_clicked:
        clear_review_decision_widgets(st.session_state.get("review"))
        plain = strip_voicevox_ruby(st.session_state.raw_script)
        st.session_state.review = None
        st.session_state.review_done = True
        st.session_state.skip_review = True
        st.session_state.script_confirmed = False
        st.session_state.ruby_ready = False
        st.session_state.mp4_bytes = None
        st.session_state.mp4_path = ""
        st.session_state.review_apply_log = []
        st.session_state.review_manual_log = []
        st.session_state.last_error = ""
        st.session_state.final_script = plain
        st.session_state.final_script_editor = plain
        st.session_state.final_script_editor_widget = plain
        st.success("レビューをスキップしました（ルビなし原稿）")

    # ----- Step 2: レビュー結果と採否／またはスキップ後の確認 -----
    if st.session_state.review_done:
        if st.session_state.get("skip_review"):
            st.write("2. ルビなし原稿を確定")
            if "final_script_editor_widget" not in st.session_state:
                st.session_state.final_script_editor_widget = strip_voicevox_ruby(
                    st.session_state.final_script or st.session_state.raw_script
                )
            st.text_area(
                "最終台本（ルビなし）",
                height=320,
                key="final_script_editor_widget",
            )
            if st.button("2. 確定してルビ準備へ", type="primary", key="btn_confirm_skip"):
                edited = strip_voicevox_ruby(
                    (st.session_state.get("final_script_editor_widget") or "").strip()
                )
                if not edited:
                    st.error("最終台本が空です。")
                else:
                    st.session_state.final_script = edited
                    st.session_state.final_script_editor = edited
                    st.session_state.raw_script = edited
                    st.session_state.script_confirmed = True
                    st.session_state.ruby_ready = False
                    st.session_state.ruby_script = ""
                    st.session_state.pop("ruby_script_editor", None)
                    st.rerun()

        elif st.session_state.review:
            st.write("2. レビュー")
            review = st.session_state.review
            mode = review.get("mode", "claude")
            if mode == "heuristic":
                st.caption("簡易レビュー")
            if review.get("review_truncated"):
                st.warning("長い台本のため、レビューは先頭部分のみです。")

            for section_key, section_title in REVIEW_SECTION_DEFS:
                render_review_section_interactive(
                    section_key,
                    section_title,
                    review.get(section_key, []),
                )

            if st.button("採択・別案を台本に反映", type="secondary"):
                base = (
                    st.session_state.get("final_script_editor_widget")
                    or st.session_state.get("final_script")
                    or st.session_state.raw_script
                )
                new_text, applied, manual = apply_review_decisions_to_script(
                    base, review
                )
                st.session_state.final_script = new_text
                st.session_state.final_script_editor = new_text
                st.session_state.final_script_editor_widget = new_text
                st.session_state.review_apply_log = applied
                st.session_state.review_manual_log = manual
                if applied:
                    st.success(f"{len(applied)} 件を反映しました")
                else:
                    st.info("自動反映できる項目はありませんでした")
                if manual:
                    st.warning("手修正が必要な項目があります")

            if st.session_state.get("review_apply_log"):
                with st.expander("反映した内容", expanded=False):
                    for line in st.session_state.review_apply_log:
                        st.write(f"- {line}")
            if st.session_state.get("review_manual_log"):
                with st.expander("手修正が必要な内容", expanded=True):
                    for line in st.session_state.review_manual_log:
                        st.write(f"- {line}")

            if "final_script_editor_widget" not in st.session_state:
                st.session_state.final_script_editor_widget = strip_voicevox_ruby(
                    st.session_state.final_script or st.session_state.raw_script
                )
            st.text_area(
                "最終台本（ルビなし）",
                height=320,
                key="final_script_editor_widget",
            )

            if st.button(
                "2. 確定してルビ準備へ",
                type="primary",
                key="btn_confirm_review",
            ):
                edited = strip_voicevox_ruby(
                    (st.session_state.get("final_script_editor_widget") or "").strip()
                )
                if not edited:
                    st.error("最終台本が空です。")
                else:
                    st.session_state.final_script = edited
                    st.session_state.final_script_editor = edited
                    st.session_state.raw_script = edited
                    st.session_state.script_confirmed = True
                    st.session_state.ruby_ready = False
                    st.session_state.ruby_script = ""
                    st.session_state.pop("ruby_script_editor", None)
                    st.rerun()

    # ----- Step 3: ルビ準備 → 動画生成 -----
    if st.session_state.script_confirmed:
        st.write("3. 動画")
        render_video_title_input()

        st.write("ルビ準備（動画の直前）")
        st.caption("①辞書 → ②ルビ作成（またはルビなしで進む）→ ③修正 → 確定、のあと動画へ")

        # ① 辞書読み込み
        st.write("① 辞書読み込み")
        dict_file = st.file_uploader(
            "追加辞書（.tsv / .txt / .csv）",
            type=["tsv", "txt", "csv"],
            key="ruby_dict_upload",
        )
        if dict_file is not None:
            try:
                raw_dict = dict_file.getvalue().decode("utf-8", errors="replace")
                pairs = parse_ruby_dict_text(raw_dict)
                if not pairs:
                    st.warning("辞書から用語を読み取れませんでした。")
                else:
                    file_id = f"{dict_file.name}-{dict_file.size}-{len(pairs)}"
                    if st.session_state.get("_ruby_dict_file_id") != file_id:
                        st.session_state.ruby_dict_custom = pairs
                        st.session_state.ruby_dict_source_name = dict_file.name
                        st.session_state._ruby_dict_file_id = file_id
                        st.session_state.ruby_ready = False
                        st.success(
                            f"追加辞書: {dict_file.name}（{len(pairs)} 語）"
                        )
            except Exception as e:  # noqa: BLE001
                st.error(f"辞書の読込失敗: {e}")
        active_n = len(get_active_ruby_dictionary())
        src_name = st.session_state.get("ruby_dict_source_name") or (
            RUBY_DICT_PATH.name if RUBY_DICT_PATH.is_file() else "組み込み"
        )
        learned_n = len(st.session_state.get("ruby_dict_learned") or [])
        st.caption(
            f"辞書: 標準 + {src_name}（{active_n} 語"
            + (f"・修正反映 {learned_n} 語" if learned_n else "")
            + "）"
        )
        if st.session_state.get("ruby_dict_export_ready") and RUBY_DICT_EXPORT_PATH.is_file():
            st.download_button(
                "ルビ辞書.txt をダウンロード",
                data=RUBY_DICT_EXPORT_PATH.read_bytes(),
                file_name=RUBY_DICT_EXPORT_NAME,
                mime="text/plain",
                key="dl_ruby_dict_pre_video",
            )

        # ② ルビ入り原稿を作成（スキップ可）
        st.write("② ルビ入り原稿を作成")
        col_ruby_on, col_ruby_skip = st.columns(2)
        with col_ruby_on:
            apply_ruby_clicked = st.button(
                "辞書でルビを付ける",
                type="primary",
                key="btn_apply_dict_ruby_pre_video",
                use_container_width=True,
            )
        with col_ruby_skip:
            skip_ruby_clicked = st.button(
                "ルビなしで進む",
                key="btn_skip_ruby_pre_video",
                use_container_width=True,
            )

        if apply_ruby_clicked:
            plain = strip_voicevox_ruby(
                st.session_state.get("final_script") or ""
            ).strip()
            if not plain:
                st.error("ルビなし原稿が空です。先に台本を確定してください。")
            else:
                with st.spinner("辞書でルビを付けています…"):
                    ruby_script, ruby_n, _ = apply_dictionary_ruby_to_script(plain)
                st.session_state.ruby_script = ruby_script
                st.session_state.ruby_script_baseline = ruby_script
                st.session_state.pop("ruby_script_editor", None)
                st.session_state.ruby_ready = False
                st.session_state.ruby_skipped = False
                st.rerun()

        if skip_ruby_clicked:
            plain = strip_voicevox_ruby(
                st.session_state.get("final_script") or ""
            ).strip()
            if not plain:
                st.error("ルビなし原稿が空です。先に台本を確定してください。")
            else:
                st.session_state.ruby_script = plain
                st.session_state.ruby_script_baseline = plain
                st.session_state.pop("ruby_script_editor", None)
                st.session_state.ruby_ready = True
                st.session_state.ruby_skipped = True
                st.session_state.ruby_dict_last_updates = []
                st.rerun()

        # ③ ルビ入り原稿を修正（確定後も上書き可）
        st.write("③ ルビ入り原稿を修正")
        has_ruby_draft = bool(
            (st.session_state.get("ruby_script") or "").strip()
            or (st.session_state.get("ruby_script_editor") or "").strip()
        )
        if not has_ruby_draft:
            st.info("②でルビを付けたあと、ここで直して確定します。")
        else:
            if "ruby_script_editor" not in st.session_state:
                st.session_state.ruby_script_editor = (
                    st.session_state.get("ruby_script") or ""
                )
            elif st.session_state.get("ruby_ready") and not (
                st.session_state.get("ruby_script_editor") or ""
            ).strip():
                st.session_state.ruby_script_editor = (
                    st.session_state.get("ruby_script") or ""
                )

            if st.session_state.get("ruby_ready"):
                n_ruby = count_voicevox_ruby(
                    st.session_state.get("ruby_script") or ""
                )
                if st.session_state.get("ruby_skipped"):
                    st.write("確定済み（ルビなしで進行）。必要なら下の欄で追記・上書きできます。")
                else:
                    st.write(f"確定済み（ルビ {n_ruby} 件）。下の欄で上書きできます。")
                last_updates = st.session_state.get("ruby_dict_last_updates") or []
                if last_updates:
                    st.caption(
                        "辞書に反映した修正: "
                        + "、".join(f"{s}→{r}" for s, r in last_updates[:12])
                        + ("…" if len(last_updates) > 12 else "")
                    )

            st.text_area(
                "ルビ入り原稿（編集・上書き可）",
                height=320,
                key="ruby_script_editor",
            )
            confirm_label = (
                "上書きして確定"
                if st.session_state.get("ruby_ready")
                else "ルビ入り原稿を確定して動画設定へ"
            )
            if st.button(
                confirm_label,
                type="primary",
                key="btn_confirm_ruby_script",
            ):
                edited_ruby = (
                    st.session_state.get("ruby_script_editor") or ""
                ).strip()
                if not edited_ruby:
                    st.error("ルビ入り原稿が空です。")
                else:
                    edited_ruby = canonicalize_voicevox_ruby_delimiters(edited_ruby)
                    baseline = str(
                        st.session_state.get("ruby_script_baseline")
                        or st.session_state.get("ruby_script")
                        or ""
                    )
                    updates = find_ruby_dict_updates(baseline, edited_ruby)
                    applied = apply_ruby_updates_to_learned_dict(updates)
                    st.session_state.ruby_script = edited_ruby
                    # ruby_script_editor は text_area の key なので、ここでは書き換えない
                    st.session_state.ruby_ready = True
                    st.session_state.ruby_skipped = False
                    st.session_state.ruby_dict_last_updates = applied
                    st.rerun()

            if RUBY_DICT_EXPORT_PATH.is_file():
                st.download_button(
                    "ルビ辞書.txt をダウンロード",
                    data=RUBY_DICT_EXPORT_PATH.read_bytes(),
                    file_name=RUBY_DICT_EXPORT_NAME,
                    mime="text/plain",
                    key="dl_ruby_dict_after_confirm",
                )
            if st.session_state.get("ruby_ready") and st.button(
                "ルビ準備をやり直す",
                key="btn_reset_ruby_ready",
            ):
                st.session_state.ruby_ready = False
                st.session_state.ruby_skipped = False
                st.rerun()

        if not st.session_state.get("ruby_ready"):
            st.caption("①②③が終わるまで、動画生成には進めません。")
            st.stop()

        st.write("参考文献")
        ref_upload = st.file_uploader(
            "参考文献ファイル（.txt / .docx）",
            type=["txt", "docx"],
            key="reference_upload",
        )
        if ref_upload is not None:
            file_id = f"{ref_upload.name}-{ref_upload.size}"
            if st.session_state.get("_reference_file_id") != file_id:
                try:
                    loaded = load_text_from_upload(ref_upload).strip()
                    st.session_state.reference_text = loaded
                    st.session_state["_reference_file_id"] = file_id
                    save_reference_text(loaded)
                    st.success(f"読込: {ref_upload.name}")
                except Exception as e:  # noqa: BLE001
                    st.error(f"参考文献の読込失敗: {e}")

        # 旧・脚注／自由文からの移行（一度だけ。生成済みエンディング全文は取り込まない）
        if st.session_state.get("_reference_migrated") is not True:
            if not (st.session_state.get("reference_text") or "").strip():
                legacy = (st.session_state.get("footnote_text") or "").strip()
                if not legacy:
                    old = (st.session_state.get("ending_credits_text") or "").strip()
                    if old and not old.startswith("本動画は医学教育用フィクション"):
                        legacy = old
                if legacy:
                    migrated = legacy.replace("医学的参考文献", "参考文献")
                    st.session_state.reference_text = migrated
                    save_reference_text(migrated)
            st.session_state["_reference_migrated"] = True

        st.text_area(
            "参考文献",
            key="reference_text",
            height=100,
            placeholder=DEFAULT_REFERENCE_EXAMPLE,
            on_change=persist_reference_from_widget,
            label_visibility="collapsed",
        )
        ref_now = (st.session_state.get("reference_text") or "").strip()
        if ref_now:
            c_ref_txt, c_ref_docx = st.columns(2)
            with c_ref_txt:
                st.download_button(
                    "参考文献 .txt",
                    data=ref_now.encode("utf-8"),
                    file_name="reference.txt",
                    mime="text/plain",
                    key="dl_reference_txt",
                )
            with c_ref_docx:
                st.download_button(
                    "参考文献 .docx",
                    data=text_to_docx_bytes(ref_now),
                    file_name="reference.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_reference_docx",
                )

        st.write("読み上げ")
        selected_style_id = int(
            st.session_state.get("vvox_style_id", DEFAULT_SPEAKER_ID)
        )
        selected_speaker_name = str(
            st.session_state.get("vvox_speaker_name", DEFAULT_SPEAKER_NAME)
        )
        selected_style_name = str(
            st.session_state.get("vvox_style_name", DEFAULT_STYLE_NAME)
        )
        try:
            ok_vv, _ver = check_voicevox()
            if not ok_vv:
                st.warning("VOICEVOXに接続できません。起動して再読み込みしてください。")
            else:
                speakers = fetch_voicevox_speakers()
                speaker_names = [
                    str(s.get("name") or "").strip()
                    for s in speakers
                    if str(s.get("name") or "").strip()
                ]
                if not speaker_names:
                    st.error("利用できる声優がありません。")
                else:
                    default_name, default_style, _default_id = (
                        resolve_default_voice_selection(speakers)
                    )
                    cur_name = st.session_state.get("vvox_speaker_name", default_name)
                    if cur_name not in speaker_names:
                        cur_name = default_name
                    name_index = speaker_names.index(cur_name)

                    col_a, col_b = st.columns(2)
                    with col_a:
                        chosen_name = st.selectbox(
                            "声優",
                            options=speaker_names,
                            index=name_index,
                            key="vvox_speaker_select",
                        )
                    speaker_obj = next(
                        s for s in speakers if str(s.get("name") or "") == chosen_name
                    )
                    styles = talk_styles_for_speaker(speaker_obj)
                    style_labels = [
                        str(s.get("name") or f"ID:{s.get('id')}") for s in styles
                    ]
                    style_ids = [int(s["id"]) for s in styles]

                    prev_name = st.session_state.get("_vvox_prev_speaker_name")
                    if prev_name != chosen_name:
                        prefer = (
                            DEFAULT_STYLE_NAME
                            if chosen_name == DEFAULT_SPEAKER_NAME
                            else "ノーマル"
                        )
                        if prefer in style_labels:
                            st.session_state.vvox_style_select = prefer
                        elif style_labels:
                            st.session_state.vvox_style_select = style_labels[0]
                        st.session_state._vvox_prev_speaker_name = chosen_name

                    cur_style = st.session_state.get("vvox_style_select")
                    if cur_style not in style_labels:
                        prefer = (
                            DEFAULT_STYLE_NAME
                            if chosen_name == DEFAULT_SPEAKER_NAME
                            else "ノーマル"
                        )
                        if prefer in style_labels:
                            style_index = style_labels.index(prefer)
                        elif (
                            chosen_name == default_name
                            and default_style in style_labels
                        ):
                            style_index = style_labels.index(default_style)
                        else:
                            style_index = 0
                    else:
                        style_index = style_labels.index(cur_style)

                    with col_b:
                        chosen_style = st.selectbox(
                            "声調",
                            options=style_labels,
                            index=style_index,
                            key="vvox_style_select",
                        )
                    selected_style_id = style_ids[style_labels.index(chosen_style)]
                    selected_speaker_name = chosen_name
                    selected_style_name = chosen_style
                    st.session_state.vvox_speaker_name = chosen_name
                    st.session_state.vvox_style_name = chosen_style
                    st.session_state.vvox_style_id = selected_style_id
        except Exception as e:  # noqa: BLE001
            st.warning(f"声優一覧の取得に失敗: {e}")
            selected_style_id = DEFAULT_SPEAKER_ID
            selected_speaker_name = DEFAULT_SPEAKER_NAME
            selected_style_name = DEFAULT_STYLE_NAME

        st.slider(
            "読み上げ速度",
            min_value=float(VOICEVOX_SPEED_MIN),
            max_value=float(VOICEVOX_SPEED_MAX),
            step=float(VOICEVOX_SPEED_STEP),
            key="vvox_speed_scale",
            format="%.1f倍",
        )

        # ----- ④ エンディング画面の確認・修正 -----
        st.write("エンディング")
        ending_speaker = str(
            st.session_state.get("vvox_speaker_name", selected_speaker_name)
            or DEFAULT_SPEAKER_NAME
        )
        ending_ref = str(st.session_state.get("reference_text") or "")
        latest_ending = build_ending_credits_text(ending_ref, ending_speaker)
        ending_sig = f"{ending_ref.strip()}\0{ending_speaker.strip()}"
        prev_auto = str(st.session_state.get("_ending_auto_text") or "")
        current_ending = str(st.session_state.get("ending_credits_text") or "")
        # 初回／空／自動文面のまま → 現在レイアウト＋最新の参考文献・声優で pre-fill
        never_inited = st.session_state.get("_ending_prefill_sig") is None
        still_auto = current_ending.strip() == prev_auto.strip()
        if never_inited or (not current_ending.strip()) or still_auto:
            st.session_state.ending_credits_text = latest_ending
            st.session_state._ending_auto_text = latest_ending
            st.session_state._ending_prefill_sig = ending_sig

        if st.button("参考文献・声優で入れ直す", key="btn_refresh_ending"):
            st.session_state.ending_credits_text = latest_ending
            st.session_state._ending_auto_text = latest_ending
            st.session_state._ending_prefill_sig = ending_sig
            st.rerun()

        st.text_area(
            "エンディング文面",
            key="ending_credits_text",
            height=280,
        )

        if st.button("3. 動画を生成する", type="primary"):
            # 次の描画で作業専用画面にし、他ボタンを出さない
            st.session_state.video_encoding = True
            st.session_state._export_job = "pending"
            st.session_state.ruby_dict_post_video_updates = []
            st.session_state.export_progress_pct = 0
            st.session_state.export_progress_msg = ""
            st.rerun()

        mp4_path = st.session_state.get("mp4_path") or ""
        if mp4_path and Path(mp4_path).exists():
            size_mb = Path(mp4_path).stat().st_size / (1024 * 1024)
            st.write(f"完成: `{mp4_path}` （約 {size_mb:.1f} MB）")
            # 大きいMP4を毎回メモリに載せると落ちやすいので上限を設ける
            if size_mb < 180:
                st.download_button(
                    label="MP4をダウンロード",
                    data=Path(mp4_path).read_bytes(),
                    file_name=st.session_state.mp4_name,
                    mime="video/mp4",
                    type="primary",
                )
            else:
                st.info("ファイルが大きいため、Finder で開いてください。")
                st.code(mp4_path)
            frame = OUTPUT_DIR / "last_frame.png"
            if frame.exists():
                st.image(str(frame), use_container_width=True)
            ending_prev = OUTPUT_DIR / "last_ending.png"
            if ending_prev.exists():
                st.image(str(ending_prev), use_container_width=True)
            script_saved = st.session_state.get("last_script_path") or ""
            if script_saved and Path(script_saved).exists():
                script_dl_name = (
                    st.session_state.get("last_script_name")
                    or make_script_docx_filename(
                        st.session_state.get("video_title", "")
                    )
                )
                st.download_button(
                    label="台本を Word で保存",
                    data=Path(script_saved).read_bytes(),
                    file_name=script_dl_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            # 動画作成後: ルビ入り最終原稿と辞書の差分を反映済み
            post_ruby = st.session_state.get("ruby_dict_post_video_updates") or []
            st.write("ルビ辞書（動画作成後）")
            if post_ruby:
                st.success(
                    f"ルビ入り最終原稿にあって辞書になかったルビを "
                    f"{len(post_ruby)} 件、辞書へ追加しました。"
                )
                preview = "、".join(f"{s}（{r}）" for s, r in post_ruby[:15])
                if len(post_ruby) > 15:
                    preview += "…"
                st.caption(preview)
            else:
                st.caption(
                    "今回、辞書に足りないルビはありませんでした"
                    "（原稿のルビはすべて辞書にありました）。"
                )
            if RUBY_DICT_EXPORT_PATH.is_file():
                st.download_button(
                    label="ルビ辞書.txt をダウンロード",
                    data=RUBY_DICT_EXPORT_PATH.read_bytes(),
                    file_name=RUBY_DICT_EXPORT_NAME,
                    mime="text/plain",
                    key="dl_ruby_dict_after_video",
                )

            # 次の原稿 → もう一度 MP4 を作るループ
            st.write("次の動画")
            next_script = st.file_uploader(
                "次の原稿（.txt / .docx）",
                type=["txt", "docx"],
                key="next_loop_script_upload",
            )
            if st.button(
                "この原稿で次の動画を作る",
                type="primary",
                key="btn_next_video_loop",
            ):
                if next_script is None:
                    st.error("原稿ファイル（.txt または .docx）を選んでください。")
                else:
                    try:
                        loaded = load_text_from_upload(next_script).strip()
                        plain = strip_voicevox_ruby(loaded).strip()
                        if not plain:
                            st.error("原稿が空でした。")
                        else:
                            commit_loaded_script(
                                plain,
                                f"loop-{next_script.name}-{len(plain)}",
                            )
                            # レビューを飛ばしてルビ準備へ（ループ用）
                            st.session_state.review_done = True
                            st.session_state.skip_review = True
                            st.session_state.script_confirmed = True
                            st.session_state.ruby_ready = False
                            st.session_state.ruby_skipped = False
                            st.session_state.ruby_script = ""
                            st.session_state.ruby_script_baseline = ""
                            st.session_state.pop("ruby_script_editor", None)
                            st.session_state.mp4_bytes = None
                            st.session_state.mp4_path = ""
                            st.session_state.final_script = plain
                            st.session_state.final_script_editor = plain
                            st.session_state.final_script_editor_widget = plain
                            OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
                            (OUTPUT_DIR / "last_script.txt").write_text(
                                plain, encoding="utf-8"
                            )
                            st.rerun()
                    except Exception as e:  # noqa: BLE001
                        st.error(f"原稿の取り込みに失敗しました: {e}")
        elif st.session_state.mp4_bytes:
            st.download_button(
                label="MP4をダウンロード",
                data=st.session_state.mp4_bytes,
                file_name=st.session_state.mp4_name,
                mime="video/mp4",
                type="primary",
            )
            st.video(st.session_state.mp4_bytes)


if __name__ == "__main__":
    main()
